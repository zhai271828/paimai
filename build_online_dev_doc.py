from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "拍卖师法则Online_多人联机游戏开发文档_v1.docx"
FONT = "Microsoft YaHei"
HEADING_BLUE = "2E74B5"
HEADING_DARK = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
TABLE_BORDER = "A6B3C2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float] | None = None) -> None:
    table.autofit = False
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)


def style_run(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.82)
        section.bottom_margin = Inches(0.82)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
        section.header_distance = Inches(0.49)
        section.footer_distance = Inches(0.49)

    styles = doc.styles
    specs = {
        "Normal": (10.5, "000000", False, 0, 6, 1.25),
        "Title": (24, "0B2545", True, 0, 8, 1.0),
        "Subtitle": (11, "555555", False, 0, 12, 1.15),
        "Heading 1": (16, HEADING_BLUE, True, 18, 10, 1.15),
        "Heading 2": (13, HEADING_BLUE, True, 14, 7, 1.15),
        "Heading 3": (12, HEADING_DARK, True, 10, 5, 1.15),
    }
    for name, (size, color, bold, before, after, line_spacing) in specs.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = line_spacing

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        pf = style.paragraph_format
        pf.space_after = Pt(4)
        pf.line_spacing = 1.25
        pf.left_indent = Inches(0.375)
        pf.first_line_indent = Inches(-0.188)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "拍卖师法则 Online · 多人联机游戏开发文档"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        style_run(run, 8.5, False, "666666")
    footer_p = section.footer.paragraphs[0]
    footer_p.text = "v1.0 工程方案 · 机密草案"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        style_run(run, 8.5, False, "666666")


def para(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        style_run(run)
    return p


def bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        style_run(run)


def numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        style_run(run)


def add_table(doc: Document, headers: list[str], rows: list[Iterable[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = str(text)
        set_cell_shading(hdr[idx], HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    set_table_width(table, widths)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                for run in p.runs:
                    style_run(run, 8.3 if len(headers) >= 4 else 9, ri == 0)
    return table


def callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, 10, True, HEADING_DARK)
    p.add_run("\n")
    r = p.add_run(body)
    style_run(r, 9.2, False, "000000")


def section(doc: Document, title: str, intro: str | None = None) -> None:
    doc.add_heading(title, level=1)
    if intro:
        para(doc, intro)


def subsection(doc: Document, title: str, intro: str | None = None) -> None:
    doc.add_heading(title, level=2)
    if intro:
        para(doc, intro)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("拍卖师法则 Online")
    style_run(r, 28, True, "0B2545")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("多人联机游戏开发文档")
    style_run(r, 22, True, HEADING_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("v1.0 可试玩联机版工程方案")
    style_run(r, 12, False, "555555")
    doc.add_paragraph()
    add_table(doc, ["项目", "内容"], [
        ("文档用途", "面向策划、前端、后端、规则引擎、测试和项目管理的开发执行文档"),
        ("项目路径", r"E:\code\拍卖"),
        ("当前基础", "完整 DOCX 规则包 + code/ TypeScript Web MVP 原型"),
        ("当前日期", "2026-06-12"),
        ("推荐技术栈", "TypeScript / React / Vite / Fastify / Socket.IO / 纯规则 engine"),
    ], [1.55, 4.95])
    callout(
        doc,
        "执行原则",
        "本项目不推倒现有 MVP。保留 monorepo、纯规则引擎、Socket.IO 房间和玩家私有视图架构，在此基础上补齐完整内容、阶段机、拍卖模式、交易、反制、持久化和测试。",
    )
    doc.add_page_break()


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    add_cover(doc)

    section(doc, "1. 项目概述", "《拍卖师法则 Online》是一款 3-5 人实时联机拍卖博弈游戏。玩家轮流担任主持人，利用信息差、话术、锦囊、事件和交易在 10 个拍卖日内积累声望。")
    add_table(doc, ["维度", "定义", "工程含义"], [
        ("游戏定位", "多人实时桌游 / 拍卖心理博弈 / 轻策略", "以好友房为核心，不做随机匹配优先。"),
        ("核心体验", "主持人信息优势、盲盒属性、竞价心理、黑市补给、事件扰动、秘密委托", "UI 必须同时支持公开桌面和个人隐私区。"),
        ("目标平台", "Web 优先，桌面浏览器主体验，移动端辅助适配", "桌面信息密度优先，移动端采用分栏折叠。"),
        ("当前代码基础", "React/Vite 前端、Fastify + Socket.IO 服务端、纯 TypeScript engine、基础房间系统", "继续沿用，不重写技术栈。"),
    ], [1.25, 2.25, 3.0])
    bullets(doc, [
        "核心乐趣不是数值堆叠，而是主持人“推高成交价赚佣金”与“制造流拍半价自吞”的双重动机。",
        "客户端永远只渲染玩家视图，不能持有完整真相状态。",
        "第一版成功标准是能稳定跑完一局 10 天完整规则局，而不是美术完成度。",
    ])

    section(doc, "2. 产品目标与版本范围", "版本规划按“先完整可跑、再提升表现、最后扩展运营能力”的顺序推进，避免在核心规则尚未稳定时引入账号、商业化和复杂社交系统。")
    add_table(doc, ["版本", "目标", "包含内容", "不包含内容"], [
        ("v0.1 MVP", "验证联机、房间、基础拍卖和终局", "英式/暗标、基础卡牌、基础房间、简单终局", "完整内容库、交易、反制、持久化"),
        ("v1.0", "完整跑完 10 天规则局", "12 类藏品、31 属性、43 锦囊、28 事件、52 委托、9 角色、4 种拍卖、交易、终局", "商业化账号、排位、复杂动画"),
        ("v1.1", "提升观感和运营能力", "观战、回放、动画、主题美术、更多引导、统计面板", "付费系统仍可后置"),
        ("暂不做", "控制范围", "好友房、临时昵称、房间码即可", "匹配排位、原生 App、商城、长期账号成长"),
    ], [0.8, 1.55, 2.8, 1.35])
    callout(doc, "v1.0 验收一句话", "一个没参与开发的玩家，只通过浏览器加入好友房，能在系统引导下完成 10 天完整规则局，并在终局看到可信、可解释的声望结算。")

    section(doc, "3. 核心玩法规则数字化")
    add_table(doc, ["规则项", "桌游规则", "数字化方案"], [
        ("人数", "3-5 人", "房间入座限制配置为 minPlayers=3、maxPlayers=5；当前 MVP 为 4-5，需要改。"),
        ("天数", "10 个拍卖日", "GameState.day 从 1 到 10；finalScoring 在第 10 天 freeTrade 后触发。"),
        ("黑市", "第 3、6、9 天", "dayIncome 后检测 blackMarketDays；每人购买锦囊/事件合计最多 2 张。"),
        ("单日拍品", "每天 2 件藏品", "preview 抽取 2 个 ArtifactInstance；打包拍卖会消耗两件。"),
        ("主持人", "顺时针轮换，无主持人日特殊", "hostForDay 支持 3/4/5 人表；无主持人日 currentHostId 为空并使用系统拍卖。"),
        ("隐藏信息", "传闻、故事、属性、结算价分层可见", "统一通过 getPlayerView 裁剪字段，不允许客户端收到完整 GameState。"),
    ], [0.95, 2.0, 3.55])
    subsection(doc, "3.1 单日阶段流转")
    add_table(doc, ["顺序", "阶段", "主要动作", "关键校验"], [
        ("1", "dayIncome", "每位玩家独立掷骰，获得点数 x10 银元", "日志记录骰点；不能由客户端提交骰点。"),
        ("2", "blackMarket", "第 3/6/9 天购买锦囊和事件", "购买上限、事件卡持有上限、现金检查。"),
        ("3", "preview", "展示当天 2 件藏品，主持人选拍卖方式", "主持人才可设置；无主持人日系统默认英式。"),
        ("4", "cardWindow", "预展后锦囊窗口", "按卡面时机和每日使用上限过滤。"),
        ("5", "auction", "执行英式/荷兰式/暗标/打包", "主持人禁拍、现金/贷款额度、平局/流拍。"),
        ("6", "settlement", "成交、佣金、揭晓买家私密信息", "结算价和属性只给买家。"),
        ("7", "eventWindow", "使用事件卡或判定自然事件", "每天最多一个事件；不回溯当天已成交。"),
        ("8", "freeTrade", "玩家交易、卖银行、贷款/还款", "交易双方确认，资产版本不匹配则失效。"),
    ], [0.45, 1.05, 2.6, 2.4])
    subsection(doc, "3.2 信息可见性矩阵")
    add_table(doc, ["信息", "所有玩家", "主持人", "买家", "服务端"], [
        ("名称 / 类别", "可见", "可见", "可见", "完整保存"),
        ("故事背景", "普通日不可见；无主持人日可见", "可见", "成交后仍按规则可见", "完整保存"),
        ("传闻区间", "不可见，除非锦囊/事件公开", "可见", "若曾探查则可见", "完整保存"),
        ("结算价", "不可见", "不可见", "成交后可见", "完整保存"),
        ("属性 / 赝品", "不可见", "不可见", "成交后可见", "完整保存"),
        ("委托 / 手牌", "只见数量", "只见数量", "本人可见", "完整保存"),
    ], [1.2, 1.5, 1.2, 1.2, 1.4])

    section(doc, "4. 系统架构")
    add_table(doc, ["层", "目录", "职责", "关键原则"], [
        ("共享层", "packages/shared", "类型、常量、内容数据、Socket payload", "任何前后端共享字段先在这里定义。"),
        ("规则层", "packages/engine", "纯 reducer、状态机、计分、隐藏视图投影", "无网络、无数据库、副作用最少。"),
        ("服务端", "apps/server", "房间、session、Socket.IO、持久化、广播", "只接受 action，不让客户端直接写状态。"),
        ("客户端", "apps/web", "React UI、操作表单、个人视图渲染", "不保存完整真相，不推导隐藏规则。"),
    ], [0.9, 1.35, 2.35, 1.9])
    callout(doc, "架构底线", "GameState 是服务端权威状态；PlayerView 是发送给单个玩家的裁剪视图。任何 UI 需要展示的数据，都应该从 PlayerView 来，而不是从客户端重建规则。")
    subsection(doc, "4.1 状态同步")
    bullets(doc, [
        "客户端发出 action：例如 auction:bid、card:play、trade:accept。",
        "服务端校验 session、房间、阶段、玩家权限。",
        "engine.reduceGame 计算新状态。",
        "服务端持久化 snapshot 和 action log。",
        "服务端对每位玩家调用 getPlayerView，分别推送 room:update。",
    ])
    subsection(doc, "4.2 持久化与回放")
    add_table(doc, ["对象", "存储建议", "用途"], [
        ("RoomSnapshot", "SQLite JSON 字段 / PostgreSQL jsonb", "快速恢复房间当前状态。"),
        ("ActionLog", "append-only 表", "复盘、调试、作弊审计、回放。"),
        ("Session", "token + roomId + playerId", "断线重连与刷新恢复。"),
        ("ContentVersion", "规则包版本 hash", "防止新旧内容混跑。"),
    ], [1.25, 2.2, 3.05])

    section(doc, "5. 数据模型设计")
    add_table(doc, ["模型", "核心字段", "用途", "安全注意"], [
        ("GameState", "roomId, phase, day, players, artifacts, decks, activeEffects, pendingReactions, tradeOffers", "服务端完整房间真相", "绝不原样发送给客户端。"),
        ("PlayerState", "id, nickname, seat, cash, loans, roleId, hand, eventHand, missionIds, artifacts, stats", "玩家完整状态", "手牌/委托只能给本人。"),
        ("ArtifactTemplate", "id, name, category, story, rumorMin, rumorMax, attributePool", "静态藏品模板", "故事和传闻属于条件可见。"),
        ("ArtifactInstance", "templateId, trueValue, isFake, attributeIds, ownerId, revealedTo, dayAcquired", "单局藏品实例", "结算价/属性/赝品严格裁剪。"),
        ("AuctionState", "mode, artifactIds, currentBid, currentBidderId, sealedBids, dutchPrice, status", "当前拍卖", "暗标金额只服务端可见。"),
        ("TrickCard / EventCard", "id, type, timing, targetSpec, effectSpec, description", "卡牌内容与规则入口", "不要把不可见目标信息塞到 description。"),
        ("MissionCard", "id, route, conditionSpec, reputation", "秘密委托判定", "只给本人，终局再公开。"),
        ("Role", "id, skills, triggers, usageLimits", "角色技能", "主动技能需要服务端次数校验。"),
        ("ActiveEffect", "sourceId, scope, modifier, expiresAt", "持续事件/属性/卡牌效果", "过期必须由阶段推进统一清理。"),
        ("TradeOffer", "fromId, toId, artifactIds, cash, status, version", "自由交易", "确认前检查现金和资产仍有效。"),
        ("PendingReaction", "sourceAction, eligiblePlayerIds, deadline, alreadyCountered", "反制窗口", "每个效果最多反制一次。"),
    ], [1.15, 2.25, 1.8, 1.3])
    subsection(doc, "5.1 建议新增 TypeScript 类型")
    add_table(doc, ["类型", "建议定义要点"], [
        ("AuctionMode", "'english' | 'dutch' | 'sealed' | 'bundle'"),
        ("ArtifactCategory", "12 类枚举，不再使用 MVP 的 calligraphy/bronze/jewelry/curio/celebrity 五类。"),
        ("AttributeType", "'positive' | 'negative' | 'special'"),
        ("CardTiming", "preview、cardWindow、auction、settlement、eventWindow、freeTrade、finalScoring。"),
        ("EffectSpec", "声明式效果：目标、数值、持续时间、结算钩子、是否可反制。"),
        ("Visibility", "public、hostOnly、ownerOnly、selfOnly、serverOnly。"),
    ], [1.35, 5.15])

    section(doc, "6. 阶段机设计", "阶段机是本项目最重要的工程边界。每个阶段都必须明确进入条件、允许操作、退出条件和错误处理，所有玩家操作都通过阶段校验。")
    add_table(doc, ["阶段", "进入条件", "允许操作", "退出条件", "错误处理"], [
        ("lobby", "房间创建后", "加入、准备、取消准备、开始", "房主开始且人数/准备满足", "人数不足、房间已开始、昵称无效。"),
        ("setup", "开始游戏后初始化", "系统发牌、发委托、发角色", "初始化完成", "内容库不足、重复 id、seed 失败。"),
        ("dayIncome", "每日开始", "系统掷骰发钱", "自动进入黑市或预展", "禁止玩家手动提交收入。"),
        ("blackMarket", "第 3/6/9 天", "购买锦囊/事件", "主持人或系统推进", "现金不足、超购买上限、事件卡超上限。"),
        ("preview", "抽取当天藏品", "主持人选拍卖方式、玩家信息锦囊", "拍卖方式确定", "非主持人设置、底价非法。"),
        ("cardWindow", "拍卖前", "按时机使用锦囊/角色技能", "所有响应完成或主持推进", "每日次数上限、目标非法。"),
        ("auction", "拍卖开始", "出价、弃权、喊停、暗标提交", "成交或流拍", "主持人出价、现金不足、重复提交。"),
        ("settlement", "拍卖结束", "系统付款、佣金、揭晓、属性应用", "当天拍品处理完或下一件", "赢家现金不足应被前置校验。"),
        ("eventWindow", "当天拍卖全结束", "使用事件卡、自然事件判定、反制", "事件结算完成", "一天多个事件、回溯当天成交。"),
        ("freeTrade", "事件窗口后", "交易、卖银行、贷款、还款", "推进下一天或终局", "交易资产变化、现金不足。"),
        ("finalScoring", "第 10 天 freeTrade 后", "查看结算、重新开局", "游戏结束", "禁止再修改经济状态。"),
    ], [0.8, 1.15, 1.45, 1.25, 1.85])

    section(doc, "7. 拍卖系统设计")
    add_table(doc, ["模式", "流程", "关键状态", "结算"], [
        ("英式", "主持人设起拍价；非主持人公开升价；玩家可弃权", "currentBid, currentBidderId, passedPlayerIds", "只剩最高出价者时成交；无人出价可自吞。"),
        ("荷兰式", "主持人设高于传闻上限的起拍价；服务端定时降价；玩家喊停", "dutchStart, dutchCurrent, tickMs, stoppedBy", "第一个喊停者成交；降至 0 无人喊停则弃置。"),
        ("暗标", "所有非主持人秘密提交出价；服务端开标", "sealedBids, tieBreakerRound", "最高价成交；平局者追加暗标，仍平局随机。"),
        ("打包", "当天两件拍品合并，按英式竞价", "artifactIds length=2, bundle=true", "成交后拆分持有，分别生成结算价/属性。"),
    ], [0.8, 2.35, 1.65, 1.7])
    subsection(doc, "7.1 出价合法性")
    bullets(doc, [
        "出价必须为非负整数，并满足最低加价。",
        "出价不得超过当前现金 + 可立即贷款额度。",
        "主持人不能竞拍自己主持的藏品。",
        "已经弃权的玩家不能重新进入，除非卡牌明确允许。",
        "暗标金额为 0 表示放弃，不能为负数。",
    ])
    subsection(doc, "7.2 流拍、自吞、弃置")
    add_table(doc, ["场景", "处理"], [
        ("英式无人出价", "主持人可按传闻最低值 x50% 自吞；若放弃则弃置。"),
        ("打包无人出价", "主持人可半价自吞两件；两件仍独立结算。"),
        ("荷兰式无人喊停", "降至 0 后弃置，主持人不可自吞。"),
        ("暗标全员 0", "流拍弃置，主持人不可自吞。"),
        ("无主持人日流拍", "直接弃置，无人获得。"),
    ], [1.5, 5.0])

    section(doc, "8. 卡牌与效果系统")
    add_table(doc, ["系统", "分类", "实现策略"], [
        ("锦囊", "信息、竞价、现金、干扰、反制", "先建立通用 targetSpec/effectSpec，再逐张映射。"),
        ("事件", "市场、环境、信息、经济、行业、黑市、交易", "进入 activeEvents，按 day/phase 自动过期。"),
        ("自然事件", "无人使用事件卡时 20% 概率", "eventWindow 服务端掷随机，日志公开。"),
        ("属性", "增益、负面、特殊", "成交时揭晓，终局和触发点按 hook 应用。"),
        ("角色", "主动/被动技能", "主动技能走 action，被动技能走 reducer hook。"),
    ], [0.9, 1.6, 4.0])
    subsection(doc, "8.1 效果结算顺序")
    numbers(doc, [
        "校验阶段和使用者权限。",
        "锁定目标快照，创建可反制的 pending effect。",
        "向可反制玩家推送 PendingReaction。",
        "反制窗口结束后，若未被取消，写入即时结果或 ActiveEffect。",
        "记录公开日志和私有日志。",
        "重新投影每位玩家 PlayerView。",
    ])
    subsection(doc, "8.2 卡牌实现优先级")
    add_table(doc, ["优先级", "范围", "原因"], [
        ("P0", "信息类、现金类、基础竞价类", "最容易实现，能支撑核心可玩性。"),
        ("P1", "事件市场波动、经济政策、黑市规则", "影响数值和阶段，但可通过 ActiveEffect 建模。"),
        ("P2", "干扰类、反制类", "需要 pending reaction 和目标选择，复杂度更高。"),
        ("P3", "强制交易、延迟结算、多目标转移", "最容易引发争议，先做规则测试再接 UI。"),
    ], [0.9, 2.2, 3.4])

    section(doc, "9. 隐藏信息与安全")
    add_table(doc, ["风险", "禁止做法", "正确做法"], [
        ("藏品真值泄露", "把完整 ArtifactInstance 发给客户端", "getPlayerView 根据身份裁剪 trueValue、attributeIds、isFake。"),
        ("暗标泄露", "在 PlayerView 暴露 sealedBids", "只暴露 sealedSubmittedPlayerIds。"),
        ("委托泄露", "公共玩家对象含 missionIds", "只暴露 handCount/eventCount/missionCount 或终局公开结果。"),
        ("主持人越权", "客户端自己判断可见传闻", "服务端根据 currentHostId 投影 rumorMin/rumorMax/story。"),
        ("重连作弊", "重连返回完整房间状态", "用 sessionToken 恢复对应玩家专属 PlayerView。"),
    ], [1.25, 2.15, 3.1])
    bullets(doc, [
        "所有 Socket payload 都必须经过 TypeScript 类型约束和服务端校验。",
        "日志分为 publicLog 和 privateLog；公开日志不能写入隐藏字段。",
        "Playwright 测试要检查 DOM 文本中不存在其他玩家私密信息。",
    ])

    section(doc, "10. 交易、银行与贷款")
    add_table(doc, ["模块", "流程", "关键校验"], [
        ("玩家交易", "发起 offer → 对方确认 → 服务端原子交换现金和藏品", "双方在线不是必须，但资产版本、现金和所有权必须有效。"),
        ("信息承诺", "玩家可在聊天/交易描述中承诺信息，但系统不自动验证", "避免系统替玩家强制公开隐藏信息。"),
        ("卖银行", "玩家选择藏品，按规则获得现金，藏品移出游戏", "珠宝 x100%，默认 x80%，属性/事件可覆盖。"),
        ("贷款", "每天最多 1 次，借 100 即时到账", "每日次数、终局还款、事件利息变化。"),
        ("强制清算", "终局现金不足还款时失去现金并交出藏品", "玩家选择被没收藏品；超时由系统选最低价值。"),
    ], [1.0, 3.0, 2.5])
    subsection(doc, "10.1 委托统计")
    add_table(doc, ["统计字段", "来源", "支持委托"], [
        ("totalAuctionSpend", "成交付款累计", "勤俭持家等经济类委托。"),
        ("playerTradeCount", "成功交易次数", "交易达人。"),
        ("flippedItemProfitCount", "买入后高价卖出记录", "倒手生财。"),
        ("auctionWinsByMode", "成交时记录拍卖模式", "暗标高手等拍卖类委托。"),
        ("peekedOpponentMissionCount", "信息行为记录", "知己知彼、局内人。"),
        ("hostedTotalSales", "主持成交额累计", "主持之王、中场核心。"),
    ], [1.55, 2.1, 2.85])

    section(doc, "11. 终局结算")
    add_table(doc, ["项目", "公式", "实现注意"], [
        ("现金声望", "floor(cash / 50)", "事件“现金为王”可覆盖为 /40 且有上限。"),
        ("藏品价值声望", "floor(sum(finalArtifactValue) / 50)", "finalArtifactValue 由基准价、属性、事件乘算得到。"),
        ("类别收藏奖励", "同类别 2/3/4+ 件 = +2/+4/+6", "赝品不计入；每类单独计算。"),
        ("秘密委托", "满足条件获得 8-11 声望", "每张独立判定；终局公开。"),
        ("属性声望", "按属性文本", "例如意外所得、慈善捐赠等需要特殊 hook。"),
        ("贷款惩罚", "未能还款触发清算", "先扣现金，再交出藏品或扣声望。"),
        ("平局比较", "藏品总价值 > 剩余现金 > 完成委托数量", "比较值需在 FinalScore 中保留。"),
    ], [1.25, 2.35, 2.9])
    subsection(doc, "11.1 终局翻牌 UI")
    bullets(doc, [
        "先展示每位玩家基础资产：现金、藏品、贷款。",
        "逐个揭示藏品终局价值和属性修正。",
        "逐张揭示秘密委托是否完成。",
        "汇总类别收藏奖励和贷款惩罚。",
        "最后展示排名和破同分依据。",
    ])

    section(doc, "12. 前端 UX 设计")
    add_table(doc, ["界面", "主要内容", "关键交互"], [
        ("大厅页", "昵称、创建房间、4 位房间码、玩家准备", "复制房间码、断线恢复、房主开始。"),
        ("主桌面", "阶段、主持人、拍品、日志、玩家列表", "所有阶段的核心信息入口。"),
        ("预展区", "当天 2 件藏品，按身份显示不同信息", "主持人可选拍卖模式；普通玩家可打信息锦囊。"),
        ("拍卖区", "当前价格、出价按钮、弃权、喊停、暗标输入", "根据拍卖模式切换控件。"),
        ("手牌区", "锦囊、事件、角色技能、委托", "只显示当前可用操作，禁用不可用卡。"),
        ("交易弹窗", "选择藏品、现金、目标玩家、确认状态", "双方确认后执行，失败给明确原因。"),
        ("终局页", "逐项结算、排名、委托翻牌", "可展开明细，支持截图分享。"),
        ("移动端", "底部标签：桌面/手牌/玩家/日志", "不追求全桌同时展示。"),
    ], [1.0, 2.6, 2.9])
    subsection(doc, "12.1 UI 状态原则")
    bullets(doc, [
        "按钮永远显示不可用原因，而不是静默禁用。",
        "涉及隐藏信息的 UI 使用“你可见/所有人可见”标签。",
        "交易和反制必须使用明确弹窗，避免玩家错过响应窗口。",
        "日志按公开/私有分区，私有日志只给本人。", 
    ])

    section(doc, "13. 后端 Socket API")
    add_table(doc, ["类别", "事件", "payload", "说明"], [
        ("房间", "room:create / room:join / room:resume", "nickname, joinCode, sessionToken", "返回 PlayerView 和 sessionToken。"),
        ("准备", "player:ready / room:start", "ready", "房主开始前校验人数与准备。"),
        ("阶段", "phase:advance", "{}", "只允许当前阶段控制者或系统推进。"),
        ("拍卖", "auction:setMode / auction:bid / auction:pass", "mode, amount", "英式和打包使用公开出价。"),
        ("荷兰式", "auction:stopDutch", "artifactId", "第一位到达服务端的合法喊停成交。"),
        ("暗标", "auction:submitSealed / auction:submitTieBreak", "amount", "平局者进入追加暗标。"),
        ("卡牌", "card:play / event:play / reaction:respond", "cardId, target ids, response", "反制窗口通过 PendingReaction 控制。"),
        ("交易", "trade:create / trade:accept / trade:reject / trade:cancel", "offer fields", "双方确认后原子结算。"),
        ("经济", "bank:sell / loan:take / loan:repay", "artifactId, amount", "服务端校验阶段和次数。"),
    ], [0.75, 1.7, 1.55, 2.5])
    subsection(doc, "13.1 Ack 规范")
    add_table(doc, ["结果", "格式", "说明"], [
        ("成功", "{ ok: true, view }", "返回当前玩家最新 PlayerView。"),
        ("失败", "{ ok: false, error }", "错误消息可直接展示给用户。"),
        ("等待响应", "{ ok: true, view, pendingReactionId }", "用于反制、交易确认等多步动作。"),
    ], [1.0, 2.4, 3.1])

    section(doc, "14. 内容管线")
    add_table(doc, ["内容", "数量", "源文档", "目标文件"], [
        ("藏品", "12 类 x20 = 240", "藏品卡.docx", "content/artifacts.json"),
        ("属性", "31 个", "属性.docx", "content/attributes.json"),
        ("锦囊", "43 张", "锦囊.docx", "content/tricks.json"),
        ("事件", "28 张 + 2 自然事件", "事件卡.docx", "content/events.json"),
        ("委托", "52 张", "委托.docx", "content/missions.json"),
        ("角色", "9 名 / 27 技能", "角色.docx", "content/roles.json"),
    ], [1.0, 1.0, 2.0, 2.5])
    subsection(doc, "14.1 内容校验脚本")
    bullets(doc, [
        "校验编号连续：I01-I15、B01-B08、C01-C08、D01-D07、R01-R05、E01-E28、W01-W52。",
        "校验数量：藏品 240、属性 31、锦囊 43、事件 28+2、委托 52、角色 9。",
        "校验文本：不得包含未完成编辑标记、待定、建议重做、调整后可用、⚖️ 等编辑痕迹。",
        "校验数值：委托声望 8-11，聚宝 +10%，诅咒 -5，珠宝银行回收 x100%。",
        "校验枚举：每个 content id 必须能映射到 shared 类型。",
    ])
    subsection(doc, "14.2 随机与可回放")
    bullets(doc, [
        "房间创建时生成 seed，所有洗牌和 random 使用 seed + actionIndex。",
        "ActionLog 记录每个玩家动作和服务端随机结果摘要。",
        "回放时从初始 seed 和 action log 重建 GameState。",
    ])

    section(doc, "15. 测试方案")
    add_table(doc, ["测试层", "覆盖内容", "关键用例"], [
        ("engine unit", "纯规则 reducer", "阶段流转、拍卖、结算、卡牌、事件、委托。"),
        ("hidden information", "PlayerView 裁剪", "普通玩家看不到传闻/属性/暗标/委托。"),
        ("socket integration", "服务端事件", "创建、加入、准备、开局、重连、多人 action。"),
        ("browser E2E", "真实 UI", "3/4/5 人开局、四种拍卖、交易、终局。"),
        ("content validation", "内容库质量", "数量、编号、文本、枚举、数值范围。"),
        ("simulation", "完整 10 天自动局", "随机合法行动跑 100 局无崩溃。"),
        ("reconnect", "断线恢复", "刷新、断网、双开、session 失效。"),
    ], [1.25, 2.0, 3.25])
    subsection(doc, "15.1 必须新增的高价值测试")
    bullets(doc, [
        "主持人永远不能看到属性和结算价，除非通过自吞/成交成为持有者。",
        "暗标平局追加暗标时，非平局者不能再提交。",
        "打包拍卖成交后两件藏品分别结算、分别交易、分别计入类别奖励。",
        "反制类锦囊不能反制反制，且每个效果最多被反制一次。",
        "玩家交易确认前若藏品已卖银行，则交易自动失效。",
    ])

    section(doc, "16. 开发里程碑", "开发里程碑以可验收的工程能力为单位，而不是按页面或功能清单机械切分；每个阶段都必须能通过自动化测试证明进展。")
    add_table(doc, ["里程碑", "交付物", "验收标准", "主要风险"], [
        ("M1 内容数据迁移", "JSON 内容库、shared 枚举、内容校验", "所有数量和编号检查通过", "DOCX 表格抽取不稳定。"),
        ("M2 完整阶段机", "eventWindow、自然事件、无主持人日、角色初始化", "自动跑完 10 天", "阶段权限和边界复杂。"),
        ("M3 四种拍卖", "英式、荷兰式、暗标、打包完整实现", "每种拍卖 E2E 成功", "荷兰式实时 tick 和延迟。"),
        ("M4 卡牌效果", "targetSpec/effectSpec、反制窗口、持续效果", "43 锦囊核心效果可用", "卡牌之间组合爆炸。"),
        ("M5 交易与终局", "交易系统、银行、贷款、完整终局", "终局明细可解释", "委托统计遗漏。"),
        ("M6 稳定上线", "持久化、重连、房间清理、部署", "24 小时房间恢复、check 全绿", "状态迁移和线上诊断。"),
    ], [1.0, 2.0, 1.8, 1.7])

    section(doc, "17. 技术风险与解决方案", "联机桌游的主要风险来自隐藏信息、复杂效果链和实时状态同步。下面的风险表用于指导实现优先级和测试投入。")
    add_table(doc, ["风险", "表现", "解决方案"], [
        ("隐藏信息泄露", "客户端看到不该看的 trueValue、属性、暗标", "所有输出只走 getPlayerView；测试扫描每个视图。"),
        ("反制链复杂", "卡牌相互打断导致状态卡死", "PendingReaction 显式建模，限制每效果一次反制。"),
        ("断线恢复", "玩家刷新后丢身份或重复入座", "sessionToken 持久化，服务端 session 表绑定 playerId。"),
        ("状态同步冲突", "多个玩家同时 action", "服务端串行处理房间 action，必要时加房间级 mutex。"),
        ("UI 信息过载", "玩家不知道当前该做什么", "阶段主行动区 + 只显示可用操作 + 不可用原因。"),
        ("内容膨胀", "240 藏品和多卡牌难维护", "内容 JSON + 校验脚本 + 内容版本 hash。"),
        ("规则代码漂移", "DOCX 和代码数值不一致", "内容源单一化，规则变更必须更新 content tests。"),
    ], [1.2, 2.0, 3.3])

    section(doc, "18. 附录")
    subsection(doc, "18.1 当前 MVP 能力清单")
    add_table(doc, ["模块", "已有能力"], [
        ("项目结构", "TypeScript monorepo，含 web/server/shared/engine。"),
        ("联机", "Socket.IO 房间创建、加入、resume、私有频道广播。"),
        ("玩法", "基础阶段流、晨间骰子收入、黑市购买、英式/暗标、贷款、卖银行、基础终局。"),
        ("测试", "engine unit、Socket flow、浏览器 E2E 已有基础。"),
        ("UI", "大厅、桌面、玩家列表、拍卖控件、藏品板、手牌面板、日志。"),
    ], [1.25, 5.25])
    subsection(doc, "18.2 v1.0 缺口清单")
    add_table(doc, ["缺口", "优先级"], [
        ("完整 240 藏品 / 31 属性 / 43 锦囊 / 28 事件 / 52 委托 / 9 角色导入", "P0"),
        ("3 人局支持、无主持人日、完整 10 天阶段机", "P0"),
        ("荷兰式和打包拍卖", "P0"),
        ("事件窗口、自然事件、反制窗口", "P1"),
        ("玩家交易、完整委托统计、完整终局翻牌", "P1"),
        ("持久化、断线恢复、房间清理", "P1"),
        ("移动端完整适配、动画、观战、回放", "P2"),
    ], [4.9, 1.6])
    subsection(doc, "18.3 建议目录结构")
    add_table(doc, ["路径", "用途"], [
        ("content/*.json", "完整规则内容库。"),
        ("packages/shared/src/content.ts", "导出校验后的内容数据和类型。"),
        ("packages/engine/src/effects", "卡牌、事件、属性、角色效果实现。"),
        ("packages/engine/src/scoring", "终局计分和委托判定。"),
        ("apps/server/src/persistence", "SQLite/PostgreSQL 存储。"),
        ("apps/web/src/components/game", "阶段 UI、拍卖 UI、交易 UI、终局 UI。"),
        ("apps/web/src/components/cards", "锦囊/事件/角色/委托卡片组件。"),
    ], [2.6, 3.9])
    subsection(doc, "18.4 推荐开发命令")
    add_table(doc, ["命令", "用途"], [
        ("npm run dev", "同时启动 server 和 web。"),
        ("npm run build", "构建 shared、engine、server、web。"),
        ("npm run test:unit", "运行 engine 单测。"),
        ("npm run test:socket", "运行 Socket.IO 集成流。"),
        ("npm run test:e2e", "运行浏览器 E2E。"),
        ("npm run check", "构建 + 测试总入口。"),
    ], [2.2, 4.3])
    subsection(doc, "18.5 v1.0 验收 Checklist")
    bullets(doc, [
        "3/4/5 人局都能开局并进入正确主持轮换。",
        "10 天流程可完整跑完，黑市只在第 3/6/9 天出现。",
        "四种拍卖都能成交、流拍或自吞，并正确结算。",
        "每位玩家的隐藏信息没有泄露给其他玩家。",
        "43 张锦囊、28 张事件、52 张委托、9 个角色至少有核心效果或明确实现状态。",
        "终局结算每一项都能展开解释。",
        "刷新页面后能回到原房间原座位。",
        "所有测试命令通过，内容校验无错误。",
    ])

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
