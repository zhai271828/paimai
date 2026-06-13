from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "拍卖师法则Online_多人联机游戏开发文档_v2_工程施工版.docx"

# compact_reference_guide preset, with named override for Chinese body text.
FONT_LATIN = "Calibri"
FONT_EAST_ASIA = "Microsoft YaHei"
MONO_FONT = "Consolas"
HEADING_BLUE = "2E74B5"
HEADING_DARK = "1F4D78"
INK_BLUE = "0B2545"
MUTED = "555555"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
LIGHT_GRAY = "F2F4F7"
CAUTION_FILL = "FFF7D6"
RISK_FILL = "FDECEC"
SUCCESS_FILL = "EAF7EA"
TABLE_BORDER = "A6B3C2"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, mono: bool = False) -> None:
    font = MONO_FONT if mono else FONT_LATIN
    east_asia = MONO_FONT if mono else FONT_EAST_ASIA
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, size: float, bold: bool = False, color: str = "000000", mono: bool = False) -> None:
    cell.text = ""
    lines = str(text).split("\n")
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(line)
        set_run_font(r, size=size, bold=bold, color=color, mono=mono)


def set_table_geometry(table, widths_in: Sequence[float] | None = None) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))

    if widths_in:
        widths_dxa = [int(width * 1440) for width in widths_in]
        grid = tbl.tblGrid
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            tbl.insert(0, grid)
        for child in list(grid):
            grid.remove(child)
        for width in widths_dxa:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for cell, width_in, width_dxa in zip(row.cells, widths_in, widths_dxa):
                cell.width = Inches(width_in)
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(width_dxa))


def configure_styles(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    specs = {
        "Normal": (11, "000000", False, 0, 6, 1.25),
        "Title": (24, INK_BLUE, True, 0, 8, 1.0),
        "Subtitle": (11, MUTED, False, 0, 12, 1.15),
        "Heading 1": (16, HEADING_BLUE, True, 18, 10, 1.15),
        "Heading 2": (13, HEADING_BLUE, True, 14, 7, 1.15),
        "Heading 3": (12, HEADING_DARK, True, 10, 5, 1.15),
    }
    for name, (size, color, bold, before, after, line_spacing) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = line_spacing

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(10.5)
        pf = style.paragraph_format
        pf.space_after = Pt(4)
        pf.line_spacing = 1.25
        pf.left_indent = Inches(0.375)
        pf.first_line_indent = Inches(-0.188)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "拍卖师法则 Online · 工程施工版开发报告"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        set_run_font(run, 8.5, False, "666666")

    footer_p = section.footer.paragraphs[0]
    footer_p.text = "v2 工程施工版 · 服务端权威 · 可验证交付"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        set_run_font(run, 8.5, False, "666666")


def para(doc: Document, text: str = "", style: str | None = None, *, bold: bool = False, color: str | None = None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, bold=bold, color=color)
    return p


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, 10.5)


def numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    widths: Sequence[float] | None = None,
    *,
    font_size: float | None = None,
    header_fill: str = HEADER_FILL,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_text(cell, str(text), size=font_size or (8.5 if len(headers) >= 5 else 9.2), bold=True, color=INK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), size=font_size or (8.0 if len(headers) >= 5 else 8.8))
    set_table_geometry(table, widths)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    para(doc, "")


def callout(doc: Document, title: str, body: str, *, fill: str = LIGHT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    set_run_font(r, 10, True, HEADING_DARK)
    p.add_run("\n")
    r = p.add_run(body)
    set_run_font(r, 9.2, False, "000000")
    set_table_geometry(table, [6.5])
    para(doc, "")


def code_block(doc: Document, text: str, *, size: float = 8.2) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, 120, 160, 120, 160)
    set_cell_text(cell, text.strip(), size=size, mono=True)
    set_table_geometry(table, [6.5])
    para(doc, "")


def read_docx_inventory() -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    for path in sorted(ROOT.glob("*.docx")):
        if path.name == OUTPUT.name:
            continue
        try:
            doc = Document(path)
            headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading") and p.text.strip()]
            rows.append((path.name, str(len(doc.tables)), str(sum(len(t.rows) for t in doc.tables)), "；".join(headings[:3]) or "无标题层级"))
        except Exception as exc:
            rows.append((path.name, "读取失败", "读取失败", str(exc)[:80]))
    return rows


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("拍卖师法则 Online")
    set_run_font(r, 28, True, INK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("多人联机游戏开发文档")
    set_run_font(r, 22, True, HEADING_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("v2 工程施工版 · 可直接拆分开发任务")
    set_run_font(r, 12, False, MUTED)

    para(doc, "")
    add_table(
        doc,
        ["项目", "内容"],
        [
            ("文档目标", "把规则包与现有 MVP 原型转化为可开发、可测试、可上线的工程规格。"),
            ("项目路径", str(ROOT)),
            ("当前基础", "7 个规则 DOCX + code/ TypeScript Web MVP。"),
            ("技术栈", "TypeScript / React / Vite / Fastify / Socket.IO / Vitest / Playwright。"),
            ("生成日期", "2026-06-12"),
            ("设计预设", "compact_reference_guide；中文字体使用 Microsoft YaHei 作为命名覆盖。"),
        ],
        [1.55, 4.95],
    )
    callout(
        doc,
        "工程判断",
        "v1 报告已经能说明方向，但还不能算可直接施工。本 v2 把缺失的类型迁移、Socket 合约、数据校验、阶段机边界、并发规则、测试样例和 Definition of Done 写入文档，目标是让外部开发者少猜测、少返工。",
        fill=SUCCESS_FILL,
    )
    doc.add_page_break()


def add_usage(doc: Document) -> None:
    heading(doc, "0. 如何使用本报告")
    callout(
        doc,
        "阅读顺序",
        "策划先看第 1、4、5、10、21、29 章；后端先看第 6、8、12、15、16、17 章；前端先看第 9、18、19 章；测试先看第 22、24、29 章；项目管理先看第 3、23、28、30 章。",
    )
    add_table(
        doc,
        ["读者", "必须阅读", "产出动作"],
        [
            ("项目负责人", "范围、里程碑、Definition of Done、风险", "确认 v1.0 不包含账号、排位、付费、观战和回放。"),
            ("规则工程师", "类型迁移、阶段机、拍卖、效果系统、终局结算", "拆分 engine PR，先写失败测试再改 reducer。"),
            ("后端工程师", "Socket API、房间队列、持久化、断线恢复、安全", "实现 room action queue、ActionLog、RoomSnapshot。"),
            ("前端工程师", "PlayerView、UI action matrix、阶段页面、私有视图", "按阶段渲染可用操作，不从客户端推断隐藏事实。"),
            ("QA", "测试矩阵、验收样例、内容校验", "把每条验收条件转成 Vitest/Socket/E2E 用例。"),
        ],
        [1.15, 2.1, 3.25],
    )
    bullets(
        doc,
        [
            "本文档不是商业化方案，也不是美术排期。它只覆盖 v1.0 可试玩联机版。",
            "每个不确定点都以工程假设写明；如果策划规则变更，必须同步修改内容 JSON、类型定义和测试。",
            "开发时以服务端 engine 结果为唯一真相，客户端只展示 PlayerView。",
        ]
    )


def add_review_conclusion(doc: Document) -> None:
    heading(doc, "1. Karpathy 式复核结论")
    para(doc, "这一章把 v1 报告按工程可执行性重新审视，明确哪些内容需要从方向描述升级为可验证规格。")
    callout(
        doc,
        "结论",
        "不能说 v1 已经完美无瑕。它是合格蓝图，但缺少工程团队直接开工所需的精确合同。本 v2 的成功标准是：字段有类型、事件有 payload、阶段有进出条件、效果有结算顺序、测试有输入和期望输出。",
        fill=CAUTION_FILL,
    )
    add_table(
        doc,
        ["复核原则", "v1 风险", "v2 修正"],
        [
            ("不隐藏假设", "多人实时、反制、断线、暗标平局等仍有解释空间。", "把假设写成服务端规则和错误处理。"),
            ("简单优先", "容易把卡牌效果做成大量 if/else。", "统一 EffectSpec + 少量特殊 resolver，按优先级迁移。"),
            ("外科式改动", "MVP 已有纯 engine 和私有视图，不应推倒。", "保留 monorepo，只扩类型、阶段、数据和测试。"),
            ("可验证目标", "“完整联机版”太宽泛。", "每个里程碑写交付物、验收命令和失败条件。"),
        ],
        [1.15, 2.55, 2.8],
    )


def add_scope(doc: Document) -> None:
    heading(doc, "2. 产品目标与版本范围")
    add_table(
        doc,
        ["版本", "目标", "必须完成", "明确不做"],
        [
            ("v0.1 MVP", "验证实时房间、基础拍卖和隐藏信息。", "英式、暗标、基础房间、少量卡牌、终局。", "完整内容、持久化、断线恢复、反制链。"),
            ("v1.0", "3-5 人完整跑完 10 天规则局。", "12 类藏品、4 种拍卖、52 委托、43 锦囊、事件、交易、贷款、终局翻牌。", "匹配排位、付费、账号成长、原生 App、观战回放。"),
            ("v1.1", "体验增强。", "动画、音效、回放、观战、房间设置、美术皮肤。", "商业化系统仍不进入核心规则 PR。"),
        ],
        [0.85, 1.65, 2.45, 1.55],
    )
    bullets(
        doc,
        [
            "v1.0 的上线定义是好友房稳定试玩，不是公开运营。",
            "每局默认 10 天，第 3、6、9 天为黑市日。",
            "首个工程目标是规则正确和隐藏信息安全，其次才是动效。",
        ]
    )


def add_current_assets(doc: Document) -> None:
    heading(doc, "3. 当前资产与 MVP 事实")
    heading(doc, "3.1 规则文档资产", 2)
    add_table(doc, ["文档", "表格数", "总行数", "标题线索"], read_docx_inventory(), [2.3, 0.7, 0.7, 2.8], font_size=7.8)
    heading(doc, "3.2 MVP 代码事实", 2)
    add_table(
        doc,
        ["位置", "当前事实", "工程含义"],
        [
            ("code/package.json", "workspaces: apps/*, packages/*；脚本含 dev/build/test/test:socket/test:e2e/check。", "继续使用现有 monorepo 和命令体系。"),
            ("packages/shared/src/types.ts", "GamePhase 缺 eventWindow；AuctionMode 只有 english/sealed；ArtifactCategory 只有 5 类。", "M1 必须先扩 shared types，否则后续内容无法落地。"),
            ("packages/shared/src/data.ts", "20 藏品、7 锦囊、3 事件、5 委托；GAME_CONSTANTS.minPlayers = 4。", "内容迁移是 v1.0 最大缺口，人数需改为 3-5。"),
            ("packages/engine/src/engine.ts", "已有纯 reducer、setup、phase advance、英式/暗标、贷款、银行、终局、getPlayerView。", "应扩展 reducer，不另起规则服务。"),
            ("apps/server/src/rooms.ts", "内存 RoomStore；按 playerChannel 推送私有 PlayerView；无持久化。", "保留私有推送模式，新增房间队列和快照存储。"),
            ("apps/web", "已有大厅、桌面、拍卖操作、手牌和藏品面板。", "按阶段矩阵补齐缺失操作，不重做 UI 框架。"),
        ],
        [1.8, 2.7, 2.0],
        font_size=8.2,
    )
    heading(doc, "3.3 关键差距表", 2)
    add_table(
        doc,
        ["模块", "当前 MVP", "v1.0 目标", "建议 PR"],
        [
            ("人数", "minPlayers 4, maxPlayers 5", "3-5 人", "PR-M1-01 constants + lobby tests"),
            ("类别", "5 类", "12 类：字画、青铜、珠宝、瓷器、玉器、古籍、钱币、奇物、灵器、邪物、遗物、绝笔", "PR-M1-02 category migration"),
            ("藏品", "20 件", "240 件", "PR-M1-03 content import + validation"),
            ("属性", "6 个 tag", "31 属性，支持多属性或按规则指定单属性池", "PR-M1-04 property schema"),
            ("委托", "每人 1 张，5 个", "每人 2 张，W01-W52，8-11 声望", "PR-M1-05 mission engine"),
            ("拍卖", "英式、暗标", "英式、荷兰式、暗标、打包；暗标平局追加一轮", "PR-M3-01 到 PR-M3-04"),
            ("流拍", "主持人可按半价自吞", "荷兰式降至 0 无人喊停则弃置；主持人禁拍规则统一", "PR-M3-05"),
            ("交易", "银行卖出", "玩家双向确认、珠宝特例、交易统计支持委托", "PR-M5-01"),
            ("持久化", "内存", "RoomSnapshot + ActionLog + Session", "PR-M6-01"),
            ("并发", "Socket handler 直接 reduce", "room action queue 串行化，reaction window 可暂停主动作", "PR-M4-03 / PR-M6-02"),
        ],
        [0.9, 1.65, 2.45, 1.5],
        font_size=7.5,
    )


def add_success_criteria(doc: Document) -> None:
    heading(doc, "4. v1.0 可交付成功标准")
    para(doc, "v1.0 的标准不是“功能看起来能点”，而是完整局、隐藏信息、安全恢复、内容一致和测试门禁同时达标。")
    add_table(
        doc,
        ["验收项", "通过标准", "验证方式"],
        [
            ("完整局", "3、4、5 人均可从大厅开局并完成 10 天。", "engine simulation + Socket integration + Playwright E2E。"),
            ("规则覆盖", "4 种拍卖、黑市、锦囊、事件、交易、贷款、委托、终局均可执行。", "测试矩阵中对应用例全部通过。"),
            ("隐藏信息", "客户端 payload 不含非授权属性、真值、暗标金额、他人委托具体条件。", "PlayerView snapshot tests + socket payload scan。"),
            ("断线恢复", "刷新或重连后恢复同一玩家身份和私有视图。", "sessionToken resume test。"),
            ("内容一致", "JSON 内容数量和编号通过校验：240/31/43/30/52/9。", "npm run validate:content。"),
            ("可维护", "每个新增效果至少有 1 个 reducer 单测或 integration 用例。", "CI 检查 coverage manifest。"),
        ],
        [1.2, 3.0, 2.3],
    )
    callout(
        doc,
        "失败即不可上线",
        "只要出现隐藏字段泄露、反制链死锁、荷兰式双赢家、暗标金额公开、RoomSnapshot 无法恢复、终局分数不可复现中的任意一项，v1.0 不应交付试玩。",
        fill=RISK_FILL,
    )


def add_rules_digitization(doc: Document) -> None:
    heading(doc, "5. 核心规则数字化")
    heading(doc, "5.1 全局常量", 2)
    add_table(
        doc,
        ["常量", "目标值", "当前 MVP", "说明"],
        [
            ("minPlayers", "3", "4", "大厅和 startGame 校验必须同步。"),
            ("maxPlayers", "5", "5", "保持。"),
            ("maxDays", "10", "10", "保持。"),
            ("blackMarketDays", "[3, 6, 9]", "[3, 6, 9]", "保持。"),
            ("startingCash", "500", "500", "保持，后续由平衡测试调整。"),
            ("hostCommissionRate", "0.2", "0.2", "保持，所有拍卖模式成交均适用，打包只按总价计算一次。"),
            ("bankSellRate", "0.8", "0.8", "按规则继续确认珠宝特例。"),
            ("loanAmount/repayment", "100/120", "100/120", "终局先扣还款，再算现金声望。"),
            ("missionRange", "8-11 reputation", "5 个 MVP 委托", "委托 JSON 校验声望必须在 8-11。"),
        ],
        [1.35, 1.25, 1.25, 2.65],
    )
    heading(doc, "5.2 稳定 key 与中文 label", 2)
    code_block(
        doc,
        """
export type ArtifactCategory =
  | "painting"    // 字画
  | "bronze"      // 青铜
  | "jewelry"     // 珠宝
  | "porcelain"   // 瓷器
  | "jade"        // 玉器
  | "book"        // 古籍
  | "coin"        // 钱币
  | "curio"       // 奇物
  | "spirit"      // 灵器
  | "cursed"      // 邪物
  | "relic"       // 遗物
  | "manuscript"; // 绝笔

export type AuctionMode = "english" | "dutch" | "sealed" | "bundle";
""",
    )
    add_table(
        doc,
        ["概念", "数字化口径", "开发注意"],
        [
            ("本日/今日", "day 对应的完整日流程，直到 freeTrade 结束。", "不要用 round 作为业务字段。"),
            ("当前拍品", "auction.artifactIds[auction.currentArtifactIndex]，打包时为 artifactIds 全部。", "所有卡牌目标必须落在当前公开上下文或玩家持有物。"),
            ("类别收藏奖励", "同类别非赝品 2/3/4+ 件 = +2/+4/+6 声望。", "替代 MVP 的 series setRep。"),
            ("赝品", "终局价值可按属性规则计算，但不计入类别收藏奖励。", "特殊角色可改价值，不改类别奖励排除。"),
            ("主持人禁拍", "主持人不能竞拍自己主持的当前拍品或打包拍品。", "服务端校验，前端按钮禁用只是体验。"),
        ],
        [1.2, 3.0, 2.3],
    )


def add_architecture(doc: Document) -> None:
    heading(doc, "6. 系统架构")
    code_block(
        doc,
        """
code/
  packages/shared/   # 类型、常量、内容 JSON schema、Socket contract
  packages/engine/   # 纯 reducer、规则校验、PlayerView 投影、计分
  apps/server/       # Fastify + Socket.IO，房间队列、持久化、session
  apps/web/          # React/Vite，阶段 UI、私有视图、操作面板
  scripts/           # 内容导入、校验、完整局模拟、E2E helper
""",
    )
    add_table(
        doc,
        ["原则", "具体要求", "禁止事项"],
        [
            ("服务端权威", "所有金额、归属、隐藏字段、抽牌、随机、计分由 server/engine 决定。", "客户端不得提交 trueValue、tag、finalScore。"),
            ("纯规则引擎", "reduceGame(state, action, context) 返回 nextState 或 RuleError。", "engine 不访问 Socket、数据库、浏览器 API。"),
            ("私有视图", "server 对每个玩家发送 getPlayerView(room, playerId)。", "不要广播完整 GameState 后让前端过滤。"),
            ("动作可回放", "所有成功 action 写入 ActionLog，可从 snapshot + log 恢复。", "不要只保存最后一个 UI 状态。"),
            ("内容可校验", "DOCX 导出的 JSON 经过 schema 和跨表引用检查。", "不要手工把规则散落到组件里。"),
        ],
        [1.25, 3.0, 2.25],
    )
    heading(doc, "6.1 运行时数据流", 2)
    numbers(
        doc,
        [
            "客户端发 Socket event，payload 只包含玩家选择，不包含结果。",
            "server 根据 socket session 绑定 playerId，进入 room action queue。",
            "engine 校验阶段、权限、现金、目标、反制窗口，生成 nextState。",
            "server 写 ActionLog，定期或关键阶段写 RoomSnapshot。",
            "server 按玩家生成 PlayerView，推送到专属 playerChannel。",
        ]
    )


def add_data_models(doc: Document) -> None:
    heading(doc, "7. 数据模型总览")
    add_table(
        doc,
        ["模型", "用途", "持久化", "隐藏信息风险"],
        [
            ("GameState", "单局完整真相。", "RoomSnapshot.state", "不能直接发给客户端。"),
            ("PlayerState", "玩家现金、手牌、事件、委托、藏品、连接状态。", "GameState.players", "他人 hand/events/missionIds 隐藏。"),
            ("ArtifactTemplate", "内容静态定义。", "content/artifacts.json", "无真值，可公开部分字段。"),
            ("ArtifactInstance", "本局生成的真值、属性、归属、揭示记录。", "GameState.artifacts", "trueValue/properties/revealedTo 高风险。"),
            ("AuctionState", "当前拍卖模式、出价、暗标、倒计时。", "GameState.auction", "sealedBids 金额高风险。"),
            ("EffectSpec", "卡牌/事件/角色效果的结构化描述。", "content + ActiveEffect", "目标选择可能暴露隐藏条件。"),
            ("PendingReaction", "可反制效果的等待窗口。", "GameState.pendingReaction", "只能给 eligible 玩家展示。"),
            ("TradeOffer", "玩家交易提案与双向确认。", "GameState.tradeOffers", "私聊式交易仅双方可见。"),
            ("ActionLog", "成功动作的可回放记录。", "数据库", "日志需保存 actor，不保存未授权视图。"),
        ],
        [1.35, 2.1, 1.4, 1.65],
        font_size=7.8,
    )
    heading(doc, "7.1 GameState 目标字段", 2)
    code_block(
        doc,
        """
interface GameState {
  schemaVersion: 2;
  contentVersion: string;
  roomId: RoomId;
  joinCode: JoinCode;
  phase: GamePhase;
  day: number;
  maxDays: 10;
  players: PlayerState[];
  artifacts: Record<ArtifactInstanceId, ArtifactInstance>;
  decks: { artifacts: ArtifactInstanceId[]; tricks: CardId[]; events: CardId[]; missions: MissionId[] };
  discardPiles: { artifacts: ArtifactInstanceId[]; tricks: CardId[]; events: CardId[] };
  currentHostId?: PlayerId;
  todayArtifactIds: ArtifactInstanceId[];
  auction?: AuctionState;
  activeEffects: ActiveEffect[];
  pendingReaction?: PendingReaction;
  tradeOffers: TradeOffer[];
  stats: GameStats;
  rng: RngState;
  actionIndex: number;
  createdAt: number;
  updatedAt: number;
}
""",
    )
    heading(doc, "7.2 PlayerState 目标字段", 2)
    code_block(
        doc,
        """
interface PlayerState {
  id: PlayerId;
  nickname: string;
  seat: number;
  ready: boolean;
  connected: boolean;
  cash: number;
  loans: LoanRecord[];
  hand: CardId[];
  events: CardId[];
  artifacts: ArtifactInstanceId[];
  missionIds: MissionId[]; // v1.0 每人 2 张
  roleId: RoleId;
  roleCharges: Record<string, number>;
  passedAuctionIds: string[];
  blackMarketBuysToday: number;
  finalScore?: FinalScore;
}
""",
    )


def add_type_migration(doc: Document) -> None:
    heading(doc, "8. TypeScript 类型迁移规格")
    para(doc, "shared 类型是前后端共同合同。迁移顺序必须先于 engine 和 UI 扩展，否则后续 PR 会反复改接口。")
    add_table(
        doc,
        ["类型", "当前", "目标", "迁移要求"],
        [
            ("GamePhase", "无 eventWindow", "加入 eventWindow；必要时加入 reactionWindow 但主状态可用 pendingReaction 表达。", "所有 phase switch 必须 exhaustive。"),
            ("AuctionMode", "english | sealed", "english | dutch | sealed | bundle", "host:setAuction payload 需支持 bundle 内部模式。"),
            ("ArtifactCategory", "5 类", "12 类稳定 key", "迁移 CATEGORY_LABELS，旧 celebrity 不再使用。"),
            ("ArtifactTag/Property", "6 个 tag", "PropertyId + properties: PropertyId[] 或单 propertyId", "推荐 properties 数组，兼容多属性角色。"),
            ("Mission", "missionId?: MissionId", "missionIds: MissionId[]", "setup 发 2 张，终局逐张判定。"),
            ("TrickCard", "type/timing/description", "category/timings/effects/counterable/target", "文本是展示，EffectSpec 是结算。"),
            ("ActiveEffect", "finalValue multiplier", "EffectSpec runtime copy + expiresAt + stackRule", "持续效果统一过期。"),
            ("Client events", "有限 MVP 事件", "完整 Socket contract", "shared 中类型即 API 合同。"),
        ],
        [1.3, 1.25, 2.25, 1.7],
        font_size=7.8,
    )
    code_block(
        doc,
        """
export type GamePhase =
  | "lobby"
  | "setup"
  | "dayIncome"
  | "blackMarket"
  | "preview"
  | "cardWindow"
  | "auction"
  | "settlement"
  | "eventWindow"
  | "freeTrade"
  | "finalScoring";

export type Visibility = "public" | "owner" | "host" | "eligibleReaction" | "finalOnly";
""",
    )


def add_content_schema(doc: Document) -> None:
    heading(doc, "9. 内容 JSON Schema")
    para(doc, "内容从 DOCX 规则表迁移到 JSON 后，规则文本只作为来源说明，运行时以 JSON 为准。每次内容变更都必须通过校验脚本。")
    heading(doc, "9.1 目录与数量", 2)
    add_table(
        doc,
        ["文件", "目标数量", "关键字段", "校验"],
        [
            ("content/artifacts.json", "240", "id, name, category, story, rumorMin, rumorMax, propertyPool", "12 类各 20 件；价格区间合法。"),
            ("content/properties.json", "31", "id, name, kind, effect", "聚宝 +10%，诅咒 -5 声望，易损/热门/传世时机明确。"),
            ("content/tricks.json", "43", "id, name, category, timings, cost, target, effects", "I/B/C/D/R 编号连续。"),
            ("content/events.json", "28 + 2", "id, name, type, timing, effects", "事件编号连续，自然事件单独标记。"),
            ("content/missions.json", "52", "id W01-W52, route, condition, reputation", "无缺号；声望 8-11。"),
            ("content/roles.json", "9", "id, name, skills, charges", "主动/被动/触发式时机清楚。"),
        ],
        [1.65, 0.75, 2.25, 1.85],
        font_size=7.8,
    )
    heading(doc, "9.2 ArtifactTemplate 示例", 2)
    code_block(
        doc,
        """
{
  "id": "painting_001",
  "name": "《寒江独钓图》",
  "category": "painting",
  "story": "传为宋人马远真迹，历代藏家钤印十三方。",
  "rumorMin": 180,
  "rumorMax": 260,
  "propertyPool": ["treasure", "heirloom", "fake", "anonymous"],
  "weight": 1,
  "contentSource": "藏品卡.docx:字画"
}
""",
    )
    heading(doc, "9.3 MissionCard 示例", 2)
    code_block(
        doc,
        """
{
  "id": "W03",
  "name": "博物志",
  "route": "collection",
  "condition": {
    "op": "countDistinctOwnedArtifacts",
    "field": "category",
    "where": { "excludeProperties": ["fake"] },
    "gte": 5
  },
  "reputation": 10,
  "description": "终局时拥有至少 5 个不同类别的藏品。"
}
""",
    )
    heading(doc, "9.4 内容校验脚本", 2)
    bullets(
        doc,
        [
            "validate:content 先做 schema 校验，再做跨表引用校验。",
            "检查编号连续性：W01-W52、锦囊分类编号、事件编号。",
            "检查禁用评估语和工程占坑标记；交付文本只保留最终卡面和最终规则。",
            "检查所有 effect.target 引用的 category/property/phase/socket action 都存在。",
            "输出 contentVersion = sha256(content/*.json)，写入 GameState。",
        ]
    )


def add_effect_spec(doc: Document) -> None:
    heading(doc, "10. EffectSpec、TargetSpec 与反制")
    para(doc, "卡牌、事件、角色技能都应尽量用同一套效果规格表达。只有少数复杂效果允许写 custom resolver，但仍要声明目标、时机、是否可反制和测试用例。")
    code_block(
        doc,
        """
interface TargetSpec {
  kind: "self" | "player" | "artifact" | "auction" | "allPlayers" | "category" | "tradeOffer";
  filters?: {
    owner?: "self" | "opponent" | "any";
    phase?: GamePhase[];
    category?: ArtifactCategory[];
    property?: PropertyId[];
    currentAuctionOnly?: boolean;
  };
  min?: number;
  max?: number;
}

interface EffectSpec {
  id: string;
  timing: EffectTiming[];
  target: TargetSpec;
  operation:
    | "reveal"
    | "modifyCash"
    | "modifyBid"
    | "modifyFinalValue"
    | "forcePass"
    | "cancelAuction"
    | "drawCard"
    | "createTradeChoice"
    | "addActiveEffect"
    | "custom";
  value?: number | string | Record<string, unknown>;
  duration?: { until: "endOfDay" | "nextBlackMarket" | "finalScoring" | "afterNextSettlement" };
  stackRule: "stack" | "maxOnly" | "replaceSameSource" | "unique";
  counterable: boolean;
  resolver?: string;
}
""",
    )
    add_table(
        doc,
        ["规则", "执行口径", "测试点"],
        [
            ("反制窗口", "counterable 效果进入 PendingReaction；eligible 玩家收到私有提示。", "只有 eligible 玩家可见和响应。"),
            ("反制次数", "每个 sourceEffect 最多被反制 1 次，不能反制反制。", "第二张反制牌被拒绝。"),
            ("超时", "默认 20 秒；所有 eligible 都 pass 或超时后继续。", "超时自动按未响应处理。"),
            ("结算顺序", "费用扣除 -> 目标锁定 -> 反制窗口 -> 效果执行 -> 日志 -> 视图推送。", "被反制时费用是否退还按卡牌字段 refundOnCounter。"),
            ("持续效果", "ActiveEffect 带 expiresAt，阶段推进时统一清理。", "第 3 天效果不影响第 4 天。"),
        ],
        [1.15, 3.15, 2.2],
    )
    code_block(
        doc,
        """
interface PendingReaction {
  id: ReactionId;
  sourceActionId: ActionId;
  sourceEffectId: string;
  sourcePlayerId: PlayerId;
  eligiblePlayerIds: PlayerId[];
  respondedPlayerIds: PlayerId[];
  expiresAt: number;
  alreadyCountered: boolean;
  pausedAction: ServerAction;
}
""",
    )


def add_phase_machine(doc: Document) -> None:
    heading(doc, "11. 阶段机设计")
    para(doc, "阶段机必须把进入条件、允许操作和退出条件写死在 reducer 中，任何 UI 判断都只能作为提示。")
    add_table(
        doc,
        ["阶段", "进入条件", "允许操作", "退出条件", "错误处理"],
        [
            ("lobby", "房间创建后", "join, ready, start", "房主 start 且人数 3-5", "未准备或人数不足返回 RuleError。"),
            ("setup", "startGame 内部短阶段", "无玩家操作", "发角色、委托、手牌、事件、洗牌完成", "任何失败回滚到 lobby。"),
            ("dayIncome", "每天开始", "phase:advance", "所有玩家收入结算完成", "只允许房主或当日主持推进。"),
            ("blackMarket", "第 3/6/9 天收入后", "buy trick/event, use blackMarket timing skill", "主持人推进", "购买次数和现金不足直接拒绝。"),
            ("preview", "生成当天 2 件藏品", "主持人查看主持层信息并选择拍卖方式", "host:setAuction", "非主持人不能设置。"),
            ("cardWindow", "拍卖方式已选", "使用竞拍前锦囊/事件/技能", "任意可推进者 advance", "若 pendingReaction 存在则不可推进。"),
            ("auction", "竞拍开始", "bid/pass/stop/submit sealed/use auction timing card", "成交或流拍", "所有出价必须服务端校验现金和身份。"),
            ("settlement", "当前拍品结束", "settlement timing card, reveal owner info", "进入下一件或 eventWindow", "结算效果按队列执行。"),
            ("eventWindow", "当天拍品全部处理后", "本日事件触发、延迟效果处理", "主持人推进", "事件没有合法目标则记录无效并弃牌。"),
            ("freeTrade", "事件窗口结束", "trade offer, bank sell, loan repay/take", "进入下一天或 finalScoring", "交易必须双方确认。"),
            ("finalScoring", "第 10 天 freeTrade 后", "查看终局翻牌", "无", "禁止任何改变状态的 action。"),
        ],
        [0.9, 1.2, 1.65, 1.3, 1.45],
        font_size=6.9,
    )
    callout(
        doc,
        "阶段机实现要求",
        "所有 action reducer 必须先 assert phase，再 assert actor authority，再 assert target legality。phase:advance 不应承担隐藏的业务分支；业务结算函数应有独立单测。",
    )


def add_auction_system(doc: Document) -> None:
    heading(doc, "12. 拍卖系统设计")
    code_block(
        doc,
        """
interface AuctionState {
  id: string;
  artifactIds: ArtifactInstanceId[];
  mode: "english" | "dutch" | "sealed" | "bundle";
  bundleInnerMode?: "english" | "dutch" | "sealed";
  currentArtifactIndex: number;
  status: "choosing" | "open" | "tieBreak" | "closed";
  currentBid: number;
  currentBidderId?: PlayerId;
  minimumIncrement: number;
  passedPlayerIds: PlayerId[];
  sealedBids: Record<PlayerId, { amount: number; round: number }>;
  tieBreakPlayerIds?: PlayerId[];
  dutch?: { startPrice: number; currentPrice: number; step: number; tickMs: number; startedAt: number };
}
""",
    )
    add_table(
        doc,
        ["模式", "流程", "成交", "边界条件"],
        [
            ("英式", "主持人设起拍价；非主持人加价或 pass；只剩最高出价者。", "赢家付款，主持人得佣金，赢家获得当前拍品。", "无出价且全 pass 则流拍，按规则弃置或进入允许的自吞逻辑。"),
            ("荷兰式", "从 startPrice 按 step/tickMs 降价；任一非主持人 stop。", "第一个被服务端队列接收的 stop 玩家按 currentPrice 成交。", "降至 0 无人喊停则流拍弃置；主持人不可自吞。"),
            ("暗标", "每位非主持人提交一次金额；金额只对本人可见。", "最高价成交。", "平局者进入 tieBreak 追加一轮暗标；仍平局再使用 seed random。"),
            ("打包", "当天 2 件藏品合并为一个 lot；主持人选择 bundleInnerMode。", "成交后两件拆分归赢家持有，成交价只扣一次，佣金按总价一次计算。", "打包消耗当天 2 件拍品，不再进入第二件独立拍卖。"),
        ],
        [0.8, 2.25, 1.75, 1.7],
        font_size=7.4,
    )
    heading(doc, "12.1 出价合法性校验", 2)
    bullets(
        doc,
        [
            "actor 必须在房间中、已连接或允许离线托管操作，且不是当前主持人。",
            "amount 必须为整数银元，不能小于 0，不能超过 actor.cash 加允许的即时贷款额度。",
            "英式 amount >= currentBid + minimumIncrement。",
            "暗标提交后可在所有人提交前修改一次，若允许修改必须写入 action log 版本。",
            "荷兰式 stop 使用服务端收到动作时的 currentPrice；相同毫秒由 room action queue 顺序决定。",
        ]
    )
    heading(doc, "12.2 sealed tie-break", 2)
    numbers(
        doc,
        [
            "第一轮暗标收齐后找出最高金额。",
            "若只有一名最高者，立即成交。",
            "若多人并列最高，AuctionState.status = tieBreak，tieBreakPlayerIds = 并列者。",
            "并列者追加一轮暗标，最低金额必须大于上一轮并列最高价；不加价可提交 0 退出。",
            "追加后仍并列，使用 GameState.rng 派生随机数选赢家，并把 seed 和候选人写入日志。",
        ]
    )


def add_cards_events_roles(doc: Document) -> None:
    heading(doc, "13. 锦囊、事件与角色系统")
    add_table(
        doc,
        ["内容类型", "运行时形态", "实现优先级", "验收"],
        [
            ("信息类锦囊", "reveal effect；写入 artifact.revealedTo 或 temporaryReveal", "P0", "不能泄露给其他玩家。"),
            ("竞价类锦囊", "modifyBid/forcePass/rejoin/cancelAuction", "P0", "阶段、当前拍品、主持人禁用规则清楚。"),
            ("现金类锦囊", "modifyCash 或 loan modifier", "P0", "现金变更进入 stats.cashDelta。"),
            ("干扰类锦囊", "target player/artifact/auction 的 ActiveEffect", "P1", "有反制窗口和过期。"),
            ("反制类锦囊", "PendingReaction response", "P0", "不能反制反制。"),
            ("事件卡", "public eventWindow 或 blackMarket 触发", "P1", "自然事件和玩家事件分开。"),
            ("角色技能", "setup 绑定 roleId；主动技能消耗 charges", "P1", "UI 显示剩余次数，服务端校验。"),
        ],
        [1.1, 2.1, 0.9, 2.4],
    )
    heading(doc, "13.1 特殊卡牌口径", 2)
    add_table(
        doc,
        ["卡/规则", "工程口径", "测试"],
        [
            ("类别嗅觉", "查看当天 2 件藏品是否同类别；打包时可同时作用于 bundle。", "返回 boolean，不返回类别名，除非卡面允许。"),
            ("来日方长", "下一件拍品成交结算后触发，不使用“下一轮”。", "若当天无下一件，则效果失效并记录。"),
            ("巧取豪夺", "创建目标玩家选择：卖给你、支付 20、展示该藏品属性。", "目标拒绝时必须二选一补偿。"),
            ("搅局类流拍", "当前拍品直接流拍弃置；主持人不可使用。", "主持人使用被拒绝。"),
            ("鉴定风潮", "赝品概率 -10%；探查赝品/属性的锦囊额外支付 10。", "内容校验标记哪些卡属于探查。"),
        ],
        [1.25, 3.35, 1.9],
        font_size=7.8,
    )


def add_hidden_info(doc: Document) -> None:
    heading(doc, "14. 隐藏信息与安全")
    add_table(
        doc,
        ["字段", "服务端真相", "本人视图", "主持人视图", "他人视图", "终局视图"],
        [
            ("artifact.trueValue", "有", "持有后可见", "不可见，除非效果允许", "不可见", "可见"),
            ("artifact.properties", "有", "持有后可见", "不可见，除非效果允许", "不可见", "可见"),
            ("artifact.rumorMin/Max", "有", "已揭示/持有可见", "预展主持层可见", "默认不可见", "可见"),
            ("sealedBids.amount", "有", "只见自己的提交状态和金额", "只见提交状态", "只见提交状态", "可在回放中选择展示"),
            ("missionIds", "有", "可见自己的 2 张", "不可见", "不可见", "终局翻牌可见"),
            ("hand/events", "有", "可见具体牌", "只见数量", "只见数量", "按产品决定是否公开"),
            ("PendingReaction", "有", "eligible 时可见", "eligible 时可见", "非 eligible 不可见", "日志可回放"),
        ],
        [1.2, 1.0, 1.25, 1.25, 1.1, 1.1],
        font_size=6.8,
    )
    heading(doc, "14.1 防作弊检查", 2)
    bullets(
        doc,
        [
            "禁止任何 ServerToClientEvents 发送 GameState。",
            "PlayerView 的 auction 字段必须 omit sealedBids 金额，只公开 sealedSubmittedPlayerIds。",
            "所有 reveal effect 只修改 revealedTo 或 PlayerView 临时字段，不改公开模板。",
            "Socket ack 返回的 view 与 room:update 使用同一 getPlayerView。",
            "测试中序列化每个玩家视图，扫描 forbidden keys: trueValue、properties、sealedBids、missionIds。终局阶段除外。",
        ]
    )


def add_socket_api(doc: Document) -> None:
    heading(doc, "15. 后端 Socket API 合同")
    para(doc, "shared/src/types.ts 是唯一 API 合同。所有事件必须 ack，ack 形态保持 { ok: true, view } 或 { ok: false, error, code? }。")
    add_table(
        doc,
        ["事件", "Payload", "权限", "成功结果", "失败码"],
        [
            ("room:create", "{ nickname }", "任意连接", "创建房间、玩家、sessionToken", "INVALID_NAME"),
            ("room:join", "{ joinCode, nickname }", "lobby 未满", "加入房间、生成 sessionToken", "ROOM_NOT_FOUND / ROOM_FULL"),
            ("room:resume", "{ roomId, playerId, sessionToken }", "session 有效", "恢复 socket data 和私有频道", "SESSION_INVALID"),
            ("player:ready", "{ ready }", "lobby 玩家", "更新 ready", "BAD_PHASE"),
            ("room:start", "{}", "房主", "setup + dayIncome", "NOT_OWNER / NOT_ENOUGH_PLAYERS"),
            ("phase:advance", "{}", "按阶段权限", "进入下一阶段", "BAD_PHASE / PENDING_REACTION"),
            ("blackMarket:buy", "{ kind }", "blackMarket 玩家", "扣钱抽牌", "LIMIT_REACHED / CASH_LOW"),
            ("host:setAuction", "{ mode, startingBid, bundleInnerMode? }", "当日主持", "进入 cardWindow", "NOT_HOST / BAD_MODE"),
            ("bid:place", "{ amount }", "英式非主持", "更新最高价", "BID_TOO_LOW / CASH_LOW"),
            ("bid:pass", "{}", "英式非主持", "退出当前拍品", "ALREADY_PASSED"),
            ("dutch:stop", "{}", "荷兰式非主持", "按 currentPrice 成交", "NOT_DUTCH / AUCTION_CLOSED"),
            ("sealedBid:submit", "{ amount }", "暗标非主持", "记录暗标或进入 tieBreak", "CASH_LOW / NOT_ELIGIBLE"),
            ("card:play", "{ cardId, targets }", "持有该卡且时机合法", "进入反制或结算效果", "CARD_NOT_OWNED / BAD_TARGET"),
            ("reaction:respond", "{ reactionId, cardId?, response }", "eligible 玩家", "反制、pass 或选择分支", "REACTION_EXPIRED"),
            ("trade:offer", "{ toPlayerId, give, receive, message? }", "freeTrade 玩家", "创建 TradeOffer", "BAD_PHASE / BAD_ASSET"),
            ("trade:respond", "{ tradeOfferId, accept, version }", "交易对方", "成交或拒绝", "VERSION_CONFLICT"),
            ("bank:sell", "{ artifactId }", "持有者", "按银行价卖出", "NOT_OWNER"),
            ("loan:take", "{}", "非终局", "获得贷款", "BAD_PHASE"),
            ("loan:repay", "{ loanId }", "借款人", "还款", "CASH_LOW"),
        ],
        [1.35, 1.75, 1.35, 1.45, 1.6],
        font_size=6.4,
    )
    code_block(
        doc,
        """
type Ack<T> =
  | ({ ok: true } & T)
  | { ok: false; error: string; code: RuleErrorCode; actionIndex?: number };

type ServerToClientEvents = {
  "room:update": (view: PlayerView) => void;
  "room:error": (payload: { message: string; code?: RuleErrorCode }) => void;
  "reaction:opened": (payload: PlayerReactionView) => void;
};
""",
    )


def add_persistence(doc: Document) -> None:
    heading(doc, "16. 持久化与日志回放")
    add_table(
        doc,
        ["表/集合", "主键", "字段", "用途"],
        [
            ("rooms", "roomId", "joinCode, status, currentSnapshotId, createdAt, updatedAt", "房间列表和恢复入口。"),
            ("room_snapshots", "snapshotId", "roomId, actionIndex, stateJson, contentVersion, createdAt", "快速恢复 GameState。"),
            ("action_logs", "roomId + actionIndex", "actionId, actorId, type, payloadJson, resultSummary, createdAt", "回放和审计。"),
            ("sessions", "sessionToken hash", "roomId, playerId, expiresAt, userAgentHash", "断线恢复。"),
            ("content_versions", "contentVersion", "sha256, counts, validationReport, createdAt", "确保旧房间用旧内容继续结算。"),
        ],
        [1.2, 1.35, 2.75, 1.2],
        font_size=7.2,
    )
    code_block(
        doc,
        """
interface RoomSnapshot {
  id: string;
  roomId: RoomId;
  actionIndex: number;
  schemaVersion: number;
  contentVersion: string;
  state: GameState;
  createdAt: number;
}

interface ActionLog {
  roomId: RoomId;
  actionIndex: number;
  actionId: ActionId;
  actorId: PlayerId;
  type: ServerAction["type"];
  payload: unknown;
  rngBefore: RngState;
  rngAfter: RngState;
  createdAt: number;
}
""",
    )
    bullets(
        doc,
        [
            "v1.0 可先使用 SQLite 或单文件 JSONL，但接口必须抽象成 RoomRepository。",
            "每个成功 action 都写 ActionLog；每个阶段结束或每 10 个 action 写 RoomSnapshot。",
            "恢复时读取最近 snapshot，再 replay 后续 action；replay 结果 actionIndex 必须一致。",
            "ContentVersion 固定在房间创建时，进行中的房间不吃新内容。",
        ]
    )


def add_concurrency(doc: Document) -> None:
    heading(doc, "17. 并发、计时与 room action queue")
    para(doc, "多人实时游戏的核心风险是同一房间内动作同时到达。本章规定服务端如何串行化动作和处理倒计时。")
    callout(
        doc,
        "核心原则",
        "每个 roomId 同一时间只能有一个 action 进入 reduceGame。Socket.IO 的并发到达顺序不能直接决定复杂规则，必须通过 room action queue 串行化并记录 actionIndex。",
    )
    code_block(
        doc,
        """
class RoomActionQueue {
  enqueue(roomId: RoomId, work: () => Promise<ActionResult>): Promise<ActionResult>;
}

// 所有 socket handler:
return queue.enqueue(roomId, async () => {
  const room = await repository.load(roomId);
  const next = reduceGame(room.state, action, { now, content, rng });
  await repository.appendActionAndMaybeSnapshot(next, action);
  broadcastPlayerViews(next);
  return { view: getPlayerView(next, actorId) };
});
""",
    )
    add_table(
        doc,
        ["场景", "规则", "测试"],
        [
            ("荷兰式同时喊停", "服务端队列中先进入者获胜；currentPrice 以该 action 处理时的 now 计算。", "两个客户端同 tick stop，只有一个赢家。"),
            ("反制窗口期间", "除 reaction:respond、room:resume 外，影响同一 sourceAction 的动作拒绝或排队。", "pendingReaction 存在时 phase:advance 失败。"),
            ("交易版本冲突", "TradeOffer 带 version；接受时版本不一致失败。", "双方同时修改/接受，只能成功一个。"),
            ("断线", "断线不自动 pass；房间设置可启用超时托管。v1.0 默认等待房主踢出或继续。", "disconnect 后 reconnect 私有视图一致。"),
            ("随机", "所有随机来自 GameState.rng，不使用 Math.random 做规则结果。", "同 seed replay 得到同一终局。"),
        ],
        [1.25, 3.1, 2.15],
    )


def add_frontend(doc: Document) -> None:
    heading(doc, "18. 前端 UX 与状态驱动界面")
    add_table(
        doc,
        ["区域", "桌面布局", "移动端策略", "数据来源"],
        [
            ("大厅", "创建/加入房间、玩家准备、房主开始。", "单列卡片，加入码输入固定 4 位。", "PlayerView.lobby fields。"),
            ("主桌面", "左侧玩家，中间预展/拍卖，右侧个人信息。", "底部 tabs：桌面/手牌/交易/日志。", "PlayerView。"),
            ("预展区", "当天 2 件藏品；主持人可见更多传闻。", "卡片横滑。", "todayArtifacts。"),
            ("拍卖操作区", "按 mode 渲染出价、pass、喊停、暗标。", "底部固定操作条。", "auction + can actions。"),
            ("手牌/事件/角色", "按 timing 高亮可用牌，目标选择弹窗。", "抽屉式面板。", "self.hand/events/role。"),
            ("交易", "提案列表、双方确认、版本冲突提示。", "逐步表单。", "tradeOffers private/public view。"),
            ("终局", "翻牌动画、分项计分、委托揭示、平局比较。", "分玩家折叠。", "finalScore。"),
        ],
        [1.05, 2.1, 1.55, 1.8],
        font_size=7.5,
    )
    heading(doc, "18.1 UI Action Matrix", 2)
    add_table(
        doc,
        ["阶段", "所有玩家", "主持人/房主", "本人私有操作"],
        [
            ("lobby", "查看玩家、复制加入码", "开始游戏", "准备/取消准备"),
            ("dayIncome", "查看收入日志", "推进", "无"),
            ("blackMarket", "查看黑市状态", "推进", "购买锦囊/事件、部分技能"),
            ("preview", "查看公开预展", "选择拍卖方式", "信息类牌若时机允许"),
            ("cardWindow", "查看当前拍品和已公开效果", "推进竞拍", "使用可用卡、选择目标"),
            ("auction", "查看竞价动态", "主持话术提示、不可竞拍", "出价/pass/喊停/暗标/竞拍时锦囊"),
            ("settlement", "查看成交", "推进", "结算时效果"),
            ("eventWindow", "查看事件", "推进", "事件响应"),
            ("freeTrade", "查看持有物数量和公开交易", "推进", "交易、银行、贷款"),
            ("finalScoring", "查看排名", "无", "翻看分项和委托"),
        ],
        [1.0, 1.8, 1.7, 2.0],
        font_size=7.4,
    )


def add_trade_bank(doc: Document) -> None:
    heading(doc, "19. 交易、银行与贷款")
    code_block(
        doc,
        """
interface TradeOffer {
  id: TradeOfferId;
  fromPlayerId: PlayerId;
  toPlayerId: PlayerId;
  give: { artifactIds: ArtifactInstanceId[]; cash: number; cardIds?: CardId[] };
  receive: { artifactIds: ArtifactInstanceId[]; cash: number; cardIds?: CardId[] };
  status: "open" | "accepted" | "rejected" | "expired" | "cancelled";
  createdAtActionIndex: number;
  expiresAtPhase: GamePhase;
  version: number;
  visibleTo: PlayerId[];
}
""",
    )
    add_table(
        doc,
        ["流程", "服务端检查", "成功后写入"],
        [
            ("创建提案", "双方存在；资产归属正确；现金充足；freeTrade 阶段。", "tradeOffers 新记录，status open。"),
            ("接受提案", "version 匹配；资产仍在双方手中；现金仍充足。", "交换资产/现金，status accepted，stats.tradeCount++。"),
            ("拒绝/取消", "操作者是 to/from。", "status rejected/cancelled。"),
            ("银行回收", "持有者拥有该藏品；阶段允许。", "现金增加，artifact.ownerId undefined 或进入 bank pile。"),
            ("珠宝特例", "若规则允许珠宝加成，按 property 或 category effect 结算。", "stats.jewelryTradeBonus。"),
            ("贷款", "非 lobby/finalScoring；未超过房间设置上限。", "LoanRecord 新增，现金 +100。"),
            ("还款", "现金 >= 120；loan 未还。", "loan.status repaid，现金 -120。"),
        ],
        [1.25, 3.15, 2.1],
    )
    bullets(
        doc,
        [
            "交易统计必须进入 GameStats，以支持委托判定，例如交易次数、卖出类别、最高成交价。",
            "银行回收不算玩家间交易，除非委托条件明确允许。",
            "强制清算仅由明确规则触发；银根紧缩只支付到 0，不触发清算。",
        ]
    )


def add_scoring(doc: Document) -> None:
    heading(doc, "20. 终局结算")
    para(doc, "终局结算必须可复现、可解释、可拆分展示。每个分项都应有独立函数和测试。")
    add_table(
        doc,
        ["分项", "公式", "实现函数", "注意"],
        [
            ("现金声望", "max(0, cash - unpaidLoanRepayment) / 50 向下取整", "scoreCash", "先处理还款和贷款惩罚。"),
            ("藏品价值声望", "sum(adjustedFinalValue) / 50 向下取整", "scoreArtifacts", "属性、事件、卡牌持续效果统一从 EffectSpec 计算。"),
            ("类别收藏奖励", "每个类别非赝品 2/3/4+ = +2/+4/+6", "scoreCategoryCollections", "替代 MVP series 奖励。"),
            ("秘密委托", "逐张判定 missionIds，成功加 8-11", "scoreMissions", "每人 2 张，失败 0。"),
            ("属性修正", "聚宝、传世、诅咒、易损等", "adjustedArtifactValue / scorePropertyRep", "声望修正和价值修正分开。"),
            ("贷款惩罚", "未能偿还的贷款按规则扣声望", "scoreLoans", "保留 loanPenalty 字段。"),
            ("平局比较", "总声望 -> 藏品价值 -> 现金 -> 最高单件价值 -> 仍平局并列", "compareWinners", "不要随机决定最终胜者。"),
        ],
        [1.25, 2.15, 1.7, 1.4],
        font_size=7.4,
    )
    code_block(
        doc,
        """
interface FinalScore {
  reputation: number;
  cashRep: number;
  artifactRep: number;
  categoryRep: number;
  missionRep: number;
  propertyRep: number;
  loanPenalty: number;
  artifactValue: number;
  tieBreakers: { artifactValue: number; cash: number; highestArtifactValue: number };
  missionResults: Array<{ missionId: MissionId; success: boolean; reputation: number }>;
}
""",
    )


def add_content_pipeline(doc: Document) -> None:
    heading(doc, "21. 内容迁移管线")
    numbers(
        doc,
        [
            "冻结规则 DOCX 为 v1.0 来源文件，复制到 content_source/。",
            "人工整理为 CSV 或直接 JSON，保留 contentSource 字段。",
            "运行 validate:content，检查数量、编号、字段、跨引用和禁用词。",
            "生成 contentVersion，并把内容作为 packages/shared 的导出。",
            "engine tests 使用固定 content fixture，E2E 使用完整 content。",
        ]
    )
    add_table(
        doc,
        ["脚本", "输入", "输出", "验收"],
        [
            ("scripts/extract-docx-tables.mjs", "规则 DOCX", "raw_content/*.csv", "能列出每张表和行号。"),
            ("scripts/normalize-content.mjs", "raw_content/*.csv", "content/*.json", "稳定 key、数字字段、效果字段规范。"),
            ("scripts/validate-content.mjs", "content/*.json", "validation-report.json", "错误数为 0。"),
            ("scripts/build-content-version.mjs", "content/*.json", "contentVersion.ts", "hash 随内容变更而变。"),
            ("scripts/simulate-full-game.mjs", "engine + content", "simulation report", "100 局无死锁、无负现金异常。"),
        ],
        [1.55, 1.4, 1.5, 2.05],
    )


def add_testing(doc: Document) -> None:
    heading(doc, "22. 测试方案")
    add_table(
        doc,
        ["层级", "工具", "覆盖", "必须样例"],
        [
            ("engine unit", "Vitest", "阶段、拍卖、效果、计分、委托", "host cannot bid; Dutch no stop discards; categoryRep excludes fake。"),
            ("content validation", "Node script + schema", "数量、编号、引用、禁用词", "W01-W52 连续，43 锦囊编号完整。"),
            ("hidden info", "Vitest snapshot", "PlayerView 字段", "他人视图无 trueValue/properties/sealedBids。"),
            ("socket integration", "socket.io-client", "多客户端动作流", "暗标 tie-break、reaction timeout、断线恢复。"),
            ("browser E2E", "Playwright", "真实 UI", "3 人完整 1 天、4 人完整 10 天 smoke。"),
            ("simulation", "engine script", "随机策略完整局", "100 局无 RuleError 死锁，终局均有分数。"),
        ],
        [1.2, 1.35, 1.65, 2.3],
        font_size=7.5,
    )
    heading(doc, "22.1 关键验收用例", 2)
    add_table(
        doc,
        ["ID", "初始状态", "动作", "期望结果"],
        [
            ("AUC-DUTCH-01", "第 1 天 auction，荷兰式 currentPrice=80，p2/p3 可喊停。", "p2 和 p3 同时 dutch:stop。", "ActionLog 顺序第一者获胜；另一 ack 失败 AUCTION_CLOSED。"),
            ("AUC-DUTCH-02", "荷兰式降到 0，无人喊停。", "timer tick 到 0。", "当前拍品流拍弃置，主持人不可自吞。"),
            ("AUC-SEALED-01", "p2/p3/p4 暗标，p2/p3 同为 90。", "收齐第一轮。", "进入 tieBreak，仅 p2/p3 可追加暗标。"),
            ("INFO-01", "p2 使用窥探查看 a01 属性。", "card:play。", "只有 p2 view 中 a01.properties 可见。"),
            ("REACTION-01", "p2 使用可反制干扰牌，p3 有反制。", "card:play。", "pendingReaction 只发给 p3，phase:advance 失败直到响应或超时。"),
            ("TRADE-01", "p2 给 p3 提案，version=1。", "p3 accept，同时 p2 cancel。", "room action queue 只允许一个成功，另一个 VERSION_CONFLICT。"),
            ("SCORE-01", "p2 持有同类别非赝品 3 件，赝品 1 件。", "finalScoring。", "categoryRep = +4，不把赝品计入 4+。"),
            ("MISSION-01", "p2 missionIds=[W01,W03]，满足 W01 不满足 W03。", "finalScoring。", "missionResults 两条，missionRep 只加 W01 声望。"),
            ("RESUME-01", "p2 持有隐藏属性和手牌。", "disconnect -> room:resume。", "恢复后 view 与断线前授权字段一致。"),
        ],
        [0.95, 2.0, 1.55, 2.0],
        font_size=6.8,
    )


def add_milestones(doc: Document) -> None:
    heading(doc, "23. 开发里程碑")
    add_table(
        doc,
        ["里程碑", "交付物", "验收标准", "主要风险"],
        [
            ("M1 内容与类型迁移", "shared types、content JSON、validate:content、3-5 人常量。", "build/test 通过；内容数量正确。", "DOCX 表格口径和 JSON 字段不一致。"),
            ("M2 完整阶段机", "eventWindow、每日 2 件、黑市、10 天完整流。", "simulation 100 局无死锁。", "阶段推进权限复杂。"),
            ("M3 四种拍卖", "English/Dutch/Sealed/Bundle 完整 reducer + UI。", "拍卖测试全通过。", "Dutch 计时并发和暗标平局。"),
            ("M4 卡牌效果与反制", "EffectSpec、PendingReaction、P0/P1 卡牌。", "每张 P0 卡至少一测。", "效果系统过度抽象或 if/else 爆炸。"),
            ("M5 交易、贷款、终局", "TradeOffer、银行、贷款、FinalScore。", "终局分数可复现。", "委托统计遗漏。"),
            ("M6 持久化与上线", "RoomRepository、ActionLog、RoomSnapshot、部署配置。", "断线恢复和 replay 通过。", "隐藏信息日志泄露。"),
        ],
        [1.1, 2.1, 1.85, 1.45],
        font_size=7.4,
    )
    heading(doc, "23.1 PR 颗粒度", 2)
    bullets(
        doc,
        [
            "每个 PR 尽量只改变一个规则面：类型、内容、阶段、拍卖、卡牌、交易、持久化分别提交。",
            "每个 PR 必须包含测试或明确说明为什么只改文档/内容。",
            "任何 PlayerView 字段变更都必须附 hidden information test。",
            "任何 GameState schema 变更都必须附 migration note 和 snapshot 兼容策略。",
        ]
    )


def add_acceptance_commands(doc: Document) -> None:
    heading(doc, "24. 开发命令与 CI Gate")
    para(doc, "以下命令构成本地合并前的最低门槛。新增模块后可以加测试，但不能降低这些 gate。")
    add_table(
        doc,
        ["命令", "用途", "通过标准"],
        [
            ("npm run build", "编译 shared/engine/server/web", "0 TypeScript error。"),
            ("npm run test", "所有 Vitest", "全部通过。"),
            ("npm run test:socket", "Socket 多客户端流程", "无未处理 rejection。"),
            ("npm run test:e2e", "Playwright 浏览器流", "关键按钮可用，无 console error。"),
            ("npm run validate:content", "内容校验", "error=0，warning 需人工确认。"),
            ("npm run simulate:full-game", "完整局模拟", "100 局完成 finalScoring。"),
            ("npm run check", "本地合并前总检查", "以上核心命令通过。"),
        ],
        [2.1, 2.4, 2.0],
    )
    callout(
        doc,
        "CI Gate",
        "main 分支必须拒绝以下情况：TypeScript 编译失败、内容校验失败、隐藏信息测试失败、任一 engine unit test 失败、Socket 基础流失败。",
        fill=RISK_FILL,
    )


def add_risks(doc: Document) -> None:
    heading(doc, "25. 技术风险与解决方案")
    para(doc, "风险管理以可观测和可测试为准：每个高风险点都要对应一个工程防线和一个检查指标。")
    add_table(
        doc,
        ["风险", "表现", "解决方案", "监控指标"],
        [
            ("隐藏信息泄露", "客户端收到 trueValue、properties、sealedBids。", "PlayerView whitelist + payload snapshot scan。", "forbidden key count = 0。"),
            ("反制链死锁", "pendingReaction 后阶段无法推进。", "超时、一次反制限制、状态机测试。", "reaction timeout recovery。"),
            ("荷兰式并发", "两个玩家同时赢。", "room action queue + actionIndex。", "每个 auction winner <= 1。"),
            ("内容漂移", "DOCX 和 JSON 不一致。", "contentSource + validation report + 内容版本。", "contentVersion diff。"),
            ("UI 信息过载", "玩家不知道该做什么。", "按 phase 只显示可用操作，日志分层。", "E2E 截图和试玩反馈。"),
            ("规则 if/else 膨胀", "卡牌难维护。", "EffectSpec 优先，custom resolver 需登记。", "custom resolver 数量。"),
            ("恢复不一致", "重连后状态跳变。", "ActionLog replay + snapshot 校验。", "replay hash match。"),
        ],
        [1.2, 1.65, 2.35, 1.3],
        font_size=7.2,
    )


def add_definition_of_done(doc: Document) -> None:
    heading(doc, "26. Definition of Done")
    para(doc, "以下条目全部满足时，才算 v1.0 工程完成；只完成 UI 或只完成规则都不能视为可交付。")
    add_table(
        doc,
        ["范围", "完成定义"],
        [
            ("规则", "规则文档中的 v1.0 必玩路径均有对应类型、reducer、UI 操作和测试。"),
            ("内容", "240 藏品、31 属性、43 锦囊、30 事件、52 委托、9 角色全部进入 JSON，并通过校验。"),
            ("联机", "3-5 人好友房能完整进行 10 天；断线可恢复；不会因普通操作卡死。"),
            ("安全", "非授权玩家永远看不到隐藏字段；暗标金额只在结算规则允许时公开。"),
            ("可回放", "RoomSnapshot + ActionLog 能恢复任意进行中房间到最新 actionIndex。"),
            ("测试", "engine、socket、hidden info、content、E2E、simulation gate 通过。"),
            ("文档", "README 写清启动、测试、内容导入和部署方式；本报告作为工程规格同步更新。"),
        ],
        [1.4, 5.1],
    )


def add_directory_plan(doc: Document) -> None:
    heading(doc, "27. 建议文件目录")
    para(doc, "目录设计保持现有 monorepo，不做大拆大建，只把规则、内容、效果、持久化边界分清。")
    code_block(
        doc,
        """
code/
  packages/shared/src/
    types.ts
    socket.ts
    constants.ts
    content/
      artifacts.json
      properties.json
      tricks.json
      events.json
      missions.json
      roles.json
    contentSchema.ts
  packages/engine/src/
    engine.ts
    phases.ts
    auctions/
      english.ts
      dutch.ts
      sealed.ts
      bundle.ts
    effects/
      resolveEffect.ts
      targetSelection.ts
      reactions.ts
    scoring.ts
    visibility.ts
  apps/server/src/
    rooms.ts
    roomActionQueue.ts
    repository/
      RoomRepository.ts
      InMemoryRoomRepository.ts
      SqliteRoomRepository.ts
  apps/web/src/
    views/
    components/auction/
    components/cards/
    components/trade/
    state/
  scripts/
    validate-content.mjs
    simulate-full-game.mjs
""",
        size=7.8,
    )


def add_task_breakdown(doc: Document) -> None:
    heading(doc, "28. PR 任务清单")
    para(doc, "任务拆分按依赖顺序排列，目标是让每个 PR 都能独立评审、独立测试、失败时容易回滚。")
    add_table(
        doc,
        ["PR", "标题", "文件", "验收"],
        [
            ("PR-M1-01", "人数和全局常量迁移", "shared/constants, engine setup, lobby UI", "3 人可开局；2 人不可开局。"),
            ("PR-M1-02", "12 类 ArtifactCategory", "shared/types, data labels, scoring", "旧 5 类测试更新；类别 label 正确。"),
            ("PR-M1-03", "导入 240 藏品", "content/artifacts.json", "12 类各 20；validate 通过。"),
            ("PR-M1-04", "31 属性模型", "properties.json, adjustedArtifactValue", "聚宝/诅咒/易损/热门/传世单测。"),
            ("PR-M1-05", "52 委托与双委托", "missions.json, scoreMissions", "每人 2 张；W01-W52 连续。"),
            ("PR-M2-01", "eventWindow 和每日流程", "engine phases, UI phase badges", "10 天模拟可结束。"),
            ("PR-M3-01", "荷兰式拍卖", "auction/dutch, socket, UI", "无人喊停弃置；同时喊停单赢家。"),
            ("PR-M3-02", "暗标 tie-break", "auction/sealed", "追加暗标和随机兜底可复现。"),
            ("PR-M3-03", "打包拍卖", "auction/bundle, scoring", "两件归属同赢家，只扣一次总价。"),
            ("PR-M4-01", "EffectSpec resolver", "effects/*", "信息/现金/价值类效果通过。"),
            ("PR-M4-02", "PendingReaction", "reactions, socket", "反制窗口和超时通过。"),
            ("PR-M5-01", "TradeOffer", "engine trade, server events, UI", "双向确认和版本冲突通过。"),
            ("PR-M5-02", "终局计分 v1.0", "scoring.ts", "类别奖励、委托、贷款、平局比较通过。"),
            ("PR-M6-01", "RoomRepository", "server repository", "snapshot + replay 恢复一致。"),
            ("PR-M6-02", "room action queue", "server queue", "并发动作顺序稳定。"),
        ],
        [0.8, 1.45, 2.1, 2.15],
        font_size=6.8,
    )


def add_appendix(doc: Document) -> None:
    heading(doc, "29. 附录：术语与工程假设")
    para(doc, "本附录统一文档中的工程术语，避免策划、前端、后端和测试使用同一个词时指向不同实现。")
    add_table(
        doc,
        ["术语", "定义"],
        [
            ("PlayerView", "某一玩家被授权看到的房间视图，是客户端唯一可信输入。"),
            ("GameState", "服务端完整真相，不直接发给客户端。"),
            ("ActionLog", "成功动作日志，用于回放、恢复和争议排查。"),
            ("ContentVersion", "内容 JSON 的 hash，锁定一局游戏使用的规则内容。"),
            ("room action queue", "按 roomId 串行化所有规则动作的队列。"),
            ("PendingReaction", "可反制效果的等待窗口，暂停源动作最终结算。"),
            ("sealed tie-break", "暗标平局追加暗标流程，仍平局再随机。"),
            ("当前拍品", "当前正在处理的单件藏品或打包 lot。"),
            ("本日", "从 dayIncome 到 freeTrade 结束的完整一天。"),
        ],
        [1.6, 4.9],
    )
    heading(doc, "30. 交付验收 Checklist")
    para(doc, "交付前按本清单逐项确认，尤其是结构检查、禁用评估语检查和渲染 QA 状态。")
    add_table(
        doc,
        ["检查项", "状态记录"],
        [
            ("文档章节不少于 25 个一级章节，包含 schema、API、任务、测试、DoD。", "交付时由结构脚本检查。"),
            ("报告不包含未完成标记或评估稿痕迹。", "交付时全文检查。"),
            ("输出 DOCX 存在且可被 python-docx 读取。", "交付时文件检查。"),
            ("若本机存在 LibreOffice，则渲染 PNG 并检查页面；若不存在，明确说明。", "交付时 QA 说明。"),
            ("后续开发以本报告 v2 为工程规格，任何规则变更同步更新文档和测试。", "项目管理确认。"),
        ],
        [4.7, 1.8],
    )


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_usage(doc)
    add_review_conclusion(doc)
    add_scope(doc)
    add_current_assets(doc)
    add_success_criteria(doc)
    add_rules_digitization(doc)
    add_architecture(doc)
    add_data_models(doc)
    add_type_migration(doc)
    add_content_schema(doc)
    add_effect_spec(doc)
    add_phase_machine(doc)
    add_auction_system(doc)
    add_cards_events_roles(doc)
    add_hidden_info(doc)
    add_socket_api(doc)
    add_persistence(doc)
    add_concurrency(doc)
    add_frontend(doc)
    add_trade_bank(doc)
    add_scoring(doc)
    add_content_pipeline(doc)
    add_testing(doc)
    add_milestones(doc)
    add_acceptance_commands(doc)
    add_risks(doc)
    add_definition_of_done(doc)
    add_directory_plan(doc)
    add_task_breakdown(doc)
    add_appendix(doc)
    doc.save(OUTPUT)


def validate_docx(path: Path) -> dict[str, object]:
    doc = Document(path)
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading") and p.text.strip()]
    h1 = [p.text.strip() for p in doc.paragraphs if p.style.name == "Heading 1" and p.text.strip()]
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    required = [
        "Socket.IO",
        "GameState",
        "PlayerView",
        "隐藏信息",
        "阶段机",
        "荷兰式",
        "打包",
        "反制",
        "交易",
        "终局结算",
        "EffectSpec",
        "TargetSpec",
        "PendingReaction",
        "TradeOffer",
        "RoomSnapshot",
        "ActionLog",
        "ContentVersion",
        "room action queue",
        "sealed tie-break",
        "content validation",
        "Definition of Done",
    ]
    banned = ["TODO", "待补", "占位", "TBD", "待定", "建议重做", "调整后可用"]
    return {
        "path": str(path),
        "exists": path.exists(),
        "size": path.stat().st_size if path.exists() else 0,
        "heading_count": len(headings),
        "h1_count": len(h1),
        "table_count": len(doc.tables),
        "missing_required": [word for word in required if word not in full_text],
        "banned_hits": [word for word in banned if word in full_text],
    }


if __name__ == "__main__":
    build_doc()
    print(json.dumps(validate_docx(OUTPUT), ensure_ascii=False, indent=2))
