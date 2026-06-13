from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "事件卡.docx"


EVENT_TYPES = [
    ("市场波动", "全局/类别", "调整后续藏品终局价值，通常为 +/-10% 到 +/-20%，持续 1 天或到下一黑市日。"),
    ("环境变化", "全局", "调整后续拍卖规则、起拍价、加价幅度或随机模式相关限制，持续 1 天。"),
    ("信息浪潮", "全局", "改变后续预展可见信息、传闻区间公开度、属性探查难度。"),
    ("经济政策", "全局", "影响贷款、银行回收价、黑市价格、晨间收入等现金流。"),
    ("行业震荡", "全局/类别", "影响赝品概率、属性费用、类别热度或保护费规则。"),
    ("黑市", "全局", "只影响下一次黑市购买上限、价格或补给结构。"),
    ("收藏/交易", "全局/条件", "鼓励玩家交易、卖银行、凑类别或调整收藏路线。"),
]


NATURAL_EVENTS = [
    {
        "id": "N1",
        "name": "神秘收购",
        "type": "自然事件",
        "timing": "后续 1 天",
        "effect": "下一个拍卖日结束后，从当天成交藏品中随机 1 件触发收购：持有者可选择以传闻最高价卖给系统；若拒绝，获得 1 声望。",
        "verdict": "调整后可用",
        "note": "原版强制买走藏品会挫败玩家；改为可选择，保留戏剧性但不强夺成果。",
    },
    {
        "id": "N2",
        "name": "经济回暖",
        "type": "自然事件",
        "timing": "后续 2 天",
        "effect": "后续 2 天内，所有玩家出售给银行的回收价提高至传闻最低值 x100%。",
        "verdict": "推荐",
        "note": "给落后玩家回血窗口，也能刺激交易与抛售，风险较低。",
    },
]


EVENT_CARDS = [
    {
        "id": "1",
        "name": "市场回暖",
        "type": "市场波动",
        "timing": "明天；1 天",
        "effect": "从明天开始，下一天所有成交藏品的终局价值 +10%。",
        "verdict": "推荐",
        "note": "简单、可预期、全员受益，适合作为基础事件。",
    },
    {
        "id": "2",
        "name": "报纸头条",
        "type": "信息浪潮",
        "timing": "明天预展",
        "effect": "明天预展时，所有人额外看到第一件藏品的故事；不公开传闻区间和属性。",
        "verdict": "推荐",
        "note": "增加谈资和判断依据，但不破坏核心信息差。",
    },
    {
        "id": "3",
        "name": "黑市打折",
        "type": "黑市",
        "timing": "下一个黑市日",
        "effect": "下一个黑市日，锦囊和事件卡价格各 -10 银元。",
        "verdict": "推荐",
        "note": "强度稳定，能制造补给期待。",
    },
    {
        "id": "4",
        "name": "黑市查封",
        "type": "黑市",
        "timing": "下一个黑市日",
        "effect": "下一个黑市日，每人最多只能买 1 张；每位玩家立刻获得 20 银元补偿。",
        "verdict": "可用",
        "note": "限制补给但给补偿，适合控节奏；注意别连续出现。",
    },
    {
        "id": "5",
        "name": "稀货流入",
        "type": "黑市",
        "timing": "下一个黑市日",
        "effect": "下一个黑市日，每人购买上限 +1；额外购买的第 3 张价格 +20 银元。",
        "verdict": "推荐",
        "note": "让现金多的人有出口，但边际成本能防止滚雪球过强。",
    },
    {
        "id": "6",
        "name": "熟人门路",
        "type": "黑市",
        "timing": "下一个黑市日",
        "effect": "下一个黑市日，每位玩家第一次购买锦囊免费；事件卡价格不变。",
        "verdict": "可用",
        "note": "趣味足，但会增加锦囊密度；建议不要和稀货流入同日叠加。",
    },
    {
        "id": "7",
        "name": "假货横行",
        "type": "行业震荡",
        "timing": "明天；1 天",
        "effect": "明天新出现藏品的赝品基础概率 +10%；若买到赝品，买家获得 30 银元补偿。",
        "verdict": "推荐",
        "note": "提高赌性并给兜底，能推动信息类锦囊价值。",
    },
    {
        "id": "8",
        "name": "鉴定风潮",
        "type": "行业震荡",
        "timing": "明天；1 天",
        "effect": "明天新出现藏品的赝品基础概率 -10%；明天所有拍品起拍价 +20 银元。",
        "verdict": "调整后可用",
        "note": "原效果只有利好，改为更安全但更贵，避免纯收益事件。",
    },
    {
        "id": "9",
        "name": "收藏展邀约",
        "type": "收藏",
        "timing": "明天结束时",
        "effect": "明天结束时，每位拥有 3 种及以上类别藏品的玩家获得 30 银元。",
        "verdict": "推荐",
        "note": "鼓励横向收藏和交易，奖励适中。",
    },
    {
        "id": "10",
        "name": "断舍离",
        "type": "交易",
        "timing": "明天自由阶段",
        "effect": "明天自由阶段，每位玩家第一次卖银行按传闻最低值 x100% 回收。",
        "verdict": "推荐",
        "note": "给现金紧张玩家解套，能减少死局。",
    },
    {
        "id": "11",
        "name": "资金冻结",
        "type": "经济政策",
        "timing": "明天；1 天",
        "effect": "明天每位玩家不能贷款；明天开始时现金低于 100 的玩家获得 40 银元补助。",
        "verdict": "推荐",
        "note": "限制爆发同时照顾落后玩家，平衡感较好。",
    },
    {
        "id": "12",
        "name": "量化宽松",
        "type": "经济政策",
        "timing": "后续 2 天",
        "effect": "后续 2 天的晨间收入阶段，所有玩家额外获得 10 银元。",
        "verdict": "可用",
        "note": "温和加钱，不会破坏经济；持续时间必须写清。",
    },
    {
        "id": "13",
        "name": "紧缩政策",
        "type": "经济政策",
        "timing": "明天；1 天",
        "effect": "明天新贷款改为借 100 还 130。已存在贷款不受影响。",
        "verdict": "可用",
        "note": "只影响新贷款，避免追溯惩罚；强度合理。",
    },
    {
        "id": "14",
        "name": "银行惜售",
        "type": "经济政策",
        "timing": "明天；1 天",
        "effect": "明天出售给银行的回收价降至传闻最低值 x70%。",
        "verdict": "调整后可用",
        "note": "原版 60% 偏狠且压制落后玩家；70% 更稳。",
    },
    {
        "id": "15",
        "name": "银行抢收",
        "type": "经济政策",
        "timing": "明天；1 天",
        "effect": "明天出售给银行的回收价提高至传闻最低值 x100%。",
        "verdict": "推荐",
        "note": "清库存、救现金，非常适合中后期。",
    },
    {
        "id": "16",
        "name": "现金为王",
        "type": "经济政策",
        "timing": "终局结算",
        "effect": "终局结算时，现金每 40 银元兑换 1 声望；由本事件额外获得的声望最多 +5。",
        "verdict": "调整后可用",
        "note": "原版每 30 银元 1 声望过强，容易让现金流路线碾压。",
    },
    {
        "id": "17",
        "name": "通胀来袭",
        "type": "经济政策",
        "timing": "明天；1 天",
        "effect": "明天黑市牌、易损保护费、贷款利息各 +10 银元；藏品成交价不受影响。",
        "verdict": "调整后可用",
        "note": "原版所有现金支付翻倍会让回合停滞；改为固定加价更好执行。",
    },
    {
        "id": "18",
        "name": "破产清算",
        "type": "经济政策",
        "timing": "明天晨间",
        "effect": "明天晨间每位玩家支付 30 银元给银行；现金不足者现金降为 0，不触发强制清算。",
        "verdict": "建议重做",
        "note": "原版立即失去 50 且可能强制清算，挫败感很强；建议只做轻惩罚。",
    },
    {
        "id": "19",
        "name": "透明市场",
        "type": "信息浪潮",
        "timing": "明天预展",
        "effect": "明天预展时，所有藏品的传闻区间对所有人公开；属性仍隐藏。",
        "verdict": "推荐",
        "note": "短暂打破主持人信息优势，能制造公平竞价日。",
    },
    {
        "id": "20",
        "name": "谣言四起",
        "type": "信息浪潮",
        "timing": "明天；1 天",
        "effect": "明天预展只公开藏品名称，不公开类别；成交后买家仍正常看到结算价和属性。",
        "verdict": "调整后可用",
        "note": "原版连买家成交后也隐藏属性，破坏反馈；保留盲拍感即可。",
    },
    {
        "id": "21",
        "name": "富豪入场",
        "type": "环境变化",
        "timing": "明天；1 天",
        "effect": "明天所有公开加价拍卖的最低加价幅度变为 30 银元；若明天随机到非公开加价模式，则第一件拍品起拍价 +20。",
        "verdict": "可用",
        "note": "能加快节奏；需兼容每日随机拍卖模式。",
    },
    {
        "id": "22",
        "name": "竞拍税",
        "type": "环境变化",
        "timing": "明天；1 天",
        "effect": "明天每位玩家第一次出价免费；之后每次加价额外支付 5 银元给银行，每人最多支付 20 银元。",
        "verdict": "调整后可用",
        "note": "原版每次出价都收费会记账繁琐；加上上限更清爽。",
    },
    {
        "id": "23",
        "name": "古籍复兴",
        "type": "市场波动",
        "timing": "明天；1 天",
        "effect": "明天【古籍/遗物/绝笔】类藏品终局价值 +15%。",
        "verdict": "推荐",
        "note": "类别组清晰，能改变次日竞价偏好。",
    },
    {
        "id": "24",
        "name": "材质危机",
        "type": "市场波动",
        "timing": "明天；1 天",
        "effect": "明天【青铜/瓷器/玉器】类藏品终局价值 -15%。",
        "verdict": "可用",
        "note": "负面类别事件可用，但建议持续 1 天，不要长期压死一类藏品。",
    },
    {
        "id": "25",
        "name": "海外热潮",
        "type": "市场波动",
        "timing": "明天；1 天",
        "effect": "明天【字画/珠宝/钱币】类藏品终局价值 +15%。",
        "verdict": "推荐",
        "note": "正向类别波动，简单好用。",
    },
    {
        "id": "26",
        "name": "灵异恐惧",
        "type": "行业震荡",
        "timing": "明天；1 天",
        "effect": "明天【灵器/邪物/奇物】类藏品持有者每件需支付 10 银元保护费；不支付则该件终局价值 -20%。",
        "verdict": "调整后可用",
        "note": "原版不支付就计 0 分过狠；-20% 保留压力但不毁局。",
    },
    {
        "id": "27",
        "name": "文化禁令",
        "type": "环境变化",
        "timing": "明天；1 天",
        "effect": "随机 1 个类别，明天该类藏品禁止玩家交易和出售给银行；终局仍可计分。",
        "verdict": "可用",
        "note": "有趣但偏限制，持续 1 天即可。",
    },
    {
        "id": "28",
        "name": "学术突破",
        "type": "行业震荡",
        "timing": "明天；1 天",
        "effect": "随机 1 个类别，明天该类新出现藏品终局价值 +20%，且赝品基础概率降为 10%。",
        "verdict": "推荐",
        "note": "强但只影响下一天新出现藏品，随机性可控。",
    },
    {
        "id": "29",
        "name": "丰收晨曦",
        "type": "经济政策",
        "timing": "明天晨间",
        "effect": "明天所有玩家晨间收入翻倍。",
        "verdict": "可用",
        "note": "由原文孤立句整理而来；效果直观，建议只持续 1 天。",
    },
]


STATUS_FILL = {
    "推荐": "D9EAD3",
    "可用": "D9EAF7",
    "调整后可用": "FFF2CC",
    "建议重做": "F4CCCC",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=9, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, size=8.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_table_borders(table, color="B7C3D0"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def apply_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, header in enumerate(headers):
        set_cell_text(hdr.cells[i], header, size=8.8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr.cells[i], "1F4E79")
        for run in hdr.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 5) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[i], value, size=font_size, align=align)
        if "推荐" in row_data:
            set_cell_shading(row.cells[5], STATUS_FILL["推荐"])
        elif "可用" in row_data:
            status = row_data[5]
            set_cell_shading(row.cells[5], STATUS_FILL.get(status, "FFFFFF"))
    apply_table_widths(table, widths)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10)


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)

    for style_name, size, color in [
        ("Heading 1", 20, "1F4E79"),
        ("Heading 2", 14, "5B7F95"),
        ("Heading 3", 11, "333333"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def main():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("事件卡")
    set_font(title_run, size=22, bold=True, color="1F4E79")
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("格式统一版 + 趣味性/平衡性评估")
    set_font(subtitle_run, size=11, color="666666")

    doc.add_heading("基本规则（统一版）", level=2)
    for line in [
        "开局每人 1 张事件卡；黑市价 50 银元/张；每人最多持有 3 张事件卡。",
        "事件在当天全部竞拍与成交结算完成后进入“竞拍后事件窗口”触发；新触发事件只影响后续，不回溯当天已成交结果。",
        "每天最多触发 1 个事件；若有玩家使用事件卡，则不再判定自然事件；若无人使用事件卡，则有 20% 概率触发自然事件。",
        "事件效果必须写明生效时间和持续时间，例如“明天；1 天”“后续 2 天”“下一个黑市日”。",
        "前几天已经生效且仍在持续的事件可以影响今天；今天竞拍后新触发的事件不能影响今天已结算藏品。",
    ]:
        add_bullet(doc, line)

    doc.add_heading("事件类型", level=2)
    add_table(
        doc,
        ["类型", "效果范围", "统一描述"],
        EVENT_TYPES,
        [1.0, 1.0, 7.8],
        font_size=9,
    )

    doc.add_heading("自然事件", level=2)
    add_table(
        doc,
        ["编号", "事件名", "类型", "生效/持续", "统一效果文本", "判定", "修改建议 / 理由"],
        [
            (item["id"], item["name"], item["type"], item["timing"], item["effect"], item["verdict"], item["note"])
            for item in NATURAL_EVENTS
        ],
        [0.45, 1.1, 0.9, 1.15, 3.7, 0.85, 1.8],
        font_size=8.4,
    )

    doc.add_heading("事件卡清单与评估", level=2)
    p = doc.add_paragraph()
    run = p.add_run("数量检查：当前文档整理出 29 张事件卡 + 2 张自然事件；原文中有重复表头、重复列和 1 条未命名效果，已合并为“丰收晨曦”。")
    set_font(run, size=9.5, color="666666")

    add_table(
        doc,
        ["#", "事件名", "类型", "生效/持续", "统一效果文本", "判定", "修改建议 / 理由"],
        [
            (item["id"], item["name"], item["type"], item["timing"], item["effect"], item["verdict"], item["note"])
            for item in EVENT_CARDS
        ],
        [0.42, 0.95, 0.82, 1.05, 3.7, 0.85, 2.0],
        font_size=7.8,
    )

    doc.add_heading("总体判断", level=2)
    for line in [
        "整体方向是好玩的：黑市、经济政策、类别涨跌、信息公开、赝品概率等都能改变玩家次日目标，适合“竞拍后影响后续”的节奏。",
        "需要重点控制的风险：不要追溯当天结果；不要直接强制清算；不要让终局现金换算、费用翻倍、藏品计 0 分这类效果无上限。",
        "最推荐保留的事件：市场回暖、报纸头条、黑市打折、稀货流入、假货横行、收藏展邀约、断舍离、资金冻结、透明市场、古籍复兴、海外热潮、学术突破。",
        "建议重做或明显削弱的事件：破产清算、通胀来袭、现金为王、谣言四起、灵异恐惧、竞拍税。它们有趣，但原版会过强、记账繁琐或破坏买家反馈。",
    ]:
        add_bullet(doc, line)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
