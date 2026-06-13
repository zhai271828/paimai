from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path
from typing import Any

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
CODE_ROOT = ROOT / "code"
CONTENT_DIR = CODE_ROOT / "content"
SHARED_CONTENT_DIR = CODE_ROOT / "packages" / "shared" / "src" / "generated"


CATEGORY_NAMES = [
    ("calligraphy", "字画"),
    ("bronze", "青铜"),
    ("jewelry", "珠宝"),
    ("porcelain", "瓷器"),
    ("jade", "玉器"),
    ("book", "古籍"),
    ("coin", "钱币"),
    ("curio", "奇物"),
    ("relic", "灵器"),
    ("evil", "邪物"),
    ("legacy", "遗物"),
    ("lastword", "绝笔"),
]


PROPERTY_ALIASES = {
    "聚宝": "treasure",
    "传世": "heirloom",
    "赝品": "fake",
    "易损": "fragile",
    "诅咒": "curse",
    "无名": "anonymous",
}

ENGINE_IMPLEMENTED_IDS = {
    "B01",
    "B02",
    "B04",
    "B05",
    "B06",
    "B08",
    "C01",
    "C03",
    "D01",
    "D02",
    "D03",
    "D04",
    "D05",
    "D06",
    "D07",
    "R01",
    "R02",
    "R03",
    "R04",
    "R05",
    "N1",
    "N2",
    "E03",
    "E05",
    "E06",
    "E10",
    "E13",
    "E14",
    "E15",
    "E16",
    "E17",
    "E18",
    "E20",
    "E21",
    "E26",
    "E28",
    "prop03",
    "prop04",
    "prop05",
    "prop08",
    "prop10",
    "prop11",
    "prop12",
    "prop13",
    "prop14",
    "prop18",
    "fragile",
    "curse",
    "prop24",
    "prop25",
    "anonymous",
    "prop31",
    "role01_skill02",
    "role01_skill03",
    "role02_skill01",
    "role02_skill02",
    "role02_skill03",
    "role03_skill01",
    "role03_skill02",
    "role03_skill03",
    "role04_skill02",
    "role04_skill03",
    "role05_skill01",
    "role05_skill02",
    "role05_skill03",
    "role06_skill02",
    "role07_skill02",
    "role07_skill03",
    "role08_skill03",
    "role09_skill01",
    "role09_skill03",
}


def clean(value: str) -> str:
    return " ".join((value or "").replace("\u3000", " ").split())


def read_tables(docx_name_part: str) -> list[list[list[str]]]:
    path = next(ROOT.glob(f"*{docx_name_part}*.docx"))
    doc = Document(str(path))
    return [[[clean(cell.text) for cell in row.cells] for row in table.rows] for table in doc.tables]


def parse_range(value: str) -> tuple[int, int]:
    numbers = [int(part) for part in re.findall(r"\d+", value)]
    if len(numbers) < 2:
        raise ValueError(f"Invalid price range: {value}")
    return numbers[0], numbers[1]


def effect_from_text(text: str, *, target: dict[str, Any] | None = None, resolver_id: str | None = None) -> list[dict[str, Any]]:
    effects: list[dict[str, Any]] = []
    normalized = clean(text)
    if match := re.search(r"终局价值\s*([+-])\s*(\d+)%", normalized):
        sign = 1 if match.group(1) == "+" else -1
        amount = int(match.group(2)) / 100
        effects.append(
            {
                "type": "finalValueMultiplier",
                "target": target or {"kind": "artifact"},
                "multiplier": round(1 + sign * amount, 3),
            }
        )
    elif match := re.search(r"价值\s*([+-])\s*(\d+)%", normalized):
        sign = 1 if match.group(1) == "+" else -1
        amount = int(match.group(2)) / 100
        effects.append(
            {
                "type": "finalValueMultiplier",
                "target": target or {"kind": "artifact"},
                "multiplier": round(1 + sign * amount, 3),
            }
        )
    if "查看" in normalized or "公开" in normalized:
        effects.append({"type": "revealInfo", "target": target or {"kind": "artifact"}, "scope": "private"})
    if money := re.search(r"获得\s*(\d+)\s*银元", normalized):
        effects.append({"type": "modifyCash", "target": {"kind": "self"}, "amount": int(money.group(1))})
    if not effects:
        resolver = f"engine:{resolver_id}" if resolver_id in ENGINE_IMPLEMENTED_IDS else "manualTextOnly"
        effects.append({"type": "custom", "target": target or {"kind": "self"}, "resolver": resolver})
    return effects


def load_properties() -> list[dict[str, Any]]:
    rows = read_tables("属性")[0][1:]
    properties: list[dict[str, Any]] = []
    for row in rows:
        if len(row) < 4 or not row[0].isdigit():
            continue
        name = row[1]
        key = PROPERTY_ALIASES.get(name, f"prop{int(row[0]):02d}")
        properties.append(
            {
                "id": key,
                "name": name,
                "kind": row[2],
                "effectText": row[3],
                "effects": effect_from_text(row[3], resolver_id=key),
                "source": {"doc": "属性.docx", "row": int(row[0])},
            }
        )
    return properties


def property_pool_for_category(index: int, all_ids: list[str]) -> list[str]:
    base = ["treasure", "heirloom", "fragile", "curse", "anonymous"]
    extra = [item for item in all_ids if item not in base and item != "fake"]
    start = (index * 2) % max(1, len(extra))
    return ["fake", *base, *extra[start : start + 4]]


def load_artifacts(property_ids: list[str]) -> list[dict[str, Any]]:
    tables = read_tables("藏品卡")
    artifacts: list[dict[str, Any]] = []
    for category_index, (category_id, category_name) in enumerate(CATEGORY_NAMES):
        rows = tables[category_index][1:]
        pool = property_pool_for_category(category_index, property_ids)
        for row in rows:
            if len(row) < 4 or not row[0].isdigit():
                continue
            rumor_min, rumor_max = parse_range(row[3])
            item_number = int(row[0])
            artifacts.append(
                {
                    "id": f"a{category_index + 1:02d}_{item_number:02d}",
                    "name": row[1].strip("《》"),
                    "category": category_id,
                    "categoryLabel": category_name,
                    "series": category_name,
                    "story": row[2],
                    "rumorMin": rumor_min,
                    "rumorMax": rumor_max,
                    "propertyPool": pool,
                    "source": {"doc": "藏品卡.docx", "category": category_name, "row": item_number},
                }
            )
    return artifacts


def load_tricks() -> list[dict[str, Any]]:
    rows = read_tables("锦囊")[1][1:]
    cards: list[dict[str, Any]] = []
    for row in rows:
        if len(row) < 7 or not row[1]:
            continue
        category = row[3]
        counterable = category != "反制类"
        cards.append(
            {
                "id": row[1],
                "name": row[2],
                "category": category,
                "type": "info" if "信息" in category else "cash" if "现金" in category else "bid" if "竞价" in category else "control",
                "timings": [row[4]],
                "cost": 0,
                "target": {"kind": "artifact" if "藏品" in row[5] or "拍品" in row[5] else "player" if "玩家" in row[5] else "self", "text": row[5]},
                "effectText": row[6],
                "effects": effect_from_text(row[6], resolver_id=row[1]),
                "counterable": counterable,
                "source": {"doc": "锦囊.docx", "row": int(row[0])},
            }
        )
    return cards


def load_events() -> list[dict[str, Any]]:
    tables = read_tables("事件卡")
    rows = [*tables[1][1:], *tables[2][1:]]
    events: list[dict[str, Any]] = []
    for row in rows:
        if len(row) < 5 or not row[0]:
            continue
        is_natural = row[0].startswith("N")
        events.append(
            {
                "id": row[0],
                "name": row[1],
                "category": row[2],
                "timing": row[3],
                "effectText": row[4],
                "effects": effect_from_text(row[4], target={"kind": "global"}, resolver_id=row[0]),
                "natural": is_natural,
                "source": {"doc": "事件卡.docx", "row": row[0]},
            }
        )
    return events


def load_missions() -> list[dict[str, Any]]:
    rows = read_tables("委托")[0][1:]
    missions: list[dict[str, Any]] = []
    for row in rows:
        if len(row) < 7 or not re.match(r"^W\d{2}$", row[0]):
            continue
        missions.append(
            {
                "id": row[0],
                "name": row[1],
                "route": row[2],
                "description": row[3],
                "condition": row[5],
                "reputation": int(re.search(r"\d+", row[4]).group(0)),
                "note": row[6],
                "source": {"doc": "委托.docx", "row": row[0]},
            }
        )
    return missions


def load_roles() -> list[dict[str, Any]]:
    path = next(ROOT.glob("*角色*.docx"))
    doc = Document(str(path))
    headings = [clean(p.text) for p in doc.paragraphs if clean(p.text).startswith("「")]
    roles: list[dict[str, Any]] = []
    for index, table in enumerate(doc.tables):
        name = headings[index].strip("「」") if index < len(headings) else f"角色{index + 1}"
        skills = []
        for row_index, row in enumerate(table.rows[1:], start=1):
            cells = [clean(cell.text) for cell in row.cells]
            if len(cells) < 4 or not cells[0]:
                continue
            skills.append(
                {
                    "id": f"role{index + 1:02d}_skill{row_index:02d}",
                    "name": cells[0],
                    "kind": cells[1],
                    "timing": cells[2],
                    "effectText": cells[3],
                    "effects": effect_from_text(cells[3], resolver_id=f"role{index + 1:02d}_skill{row_index:02d}"),
                    "charges": int(match.group(1)) if (match := re.search(r"限\s*(\d+)\s*次", cells[2])) else None,
                }
            )
        roles.append({"id": f"role{index + 1:02d}", "name": name, "skills": skills, "source": {"doc": "角色.docx", "table": index}})
    return roles


def write_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def build_version(files: dict[str, Any]) -> str:
    blob = json.dumps(files, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(blob).hexdigest()[:16]


def main() -> None:
    properties = load_properties()
    artifacts = load_artifacts([item["id"] for item in properties])
    tricks = load_tricks()
    events = load_events()
    missions = load_missions()
    roles = load_roles()
    content = {
        "artifacts": artifacts,
        "properties": properties,
        "tricks": tricks,
        "events": events,
        "missions": missions,
        "roles": roles,
    }
    content_version = build_version(content)

    for name, value in content.items():
        write_json(CONTENT_DIR / f"{name}.json", value)
    write_json(CONTENT_DIR / "content-version.json", {"contentVersion": content_version})

    SHARED_CONTENT_DIR.mkdir(parents=True, exist_ok=True)
    module = ["/* This file is generated by scripts/generate_content.py for runtime-safe ESM deployment. */"]
    for name, value in {**content, "version": {"contentVersion": content_version}}.items():
        module.append(f"const {name} = {json.dumps(value, ensure_ascii=False, separators=(',', ':'))} as const;")
    module.extend(
        [
            "",
            "export const generatedContent = { artifacts, properties, tricks, events, missions, roles, version } as const;",
            "export const CONTENT_VERSION = version.contentVersion;",
            "",
        ]
    )
    (SHARED_CONTENT_DIR / "content.ts").write_text("\n".join(module), encoding="utf-8")
    print(
        json.dumps(
            {
                "contentVersion": content_version,
                "counts": {key: len(value) for key, value in content.items()},
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
