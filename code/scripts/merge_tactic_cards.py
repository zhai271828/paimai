from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from organize_event_cards import add_bullet, configure_document, set_font
from build_tactic_cards import add_table_custom


ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "锦囊.docx"
OUT = ROOT / "锦囊.docx"


TYPE_MAP = {
    "信息类": "信息探查类",
    "竞价类": "竞拍博弈类",
    "现金类": "资金调度类",
    "交易类": "资金调度类",
    "干扰类": "行动干扰类",
    "反制类": "效果反制类",
}


SECTION_TYPE_KEYWORDS = [
    ("信息探查", "信息探查类"),
    ("竞拍博弈", "竞拍博弈类"),
    ("资金调度", "资金调度类"),
    ("行动干扰", "行动干扰类"),
    ("效果反制", "效果反制类"),
]


def iter_blocks(doc):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def clean(text):
    return " ".join((text or "").split())


def normalize_type(value):
    value = clean(value)
    return TYPE_MAP.get(value, value or "待分类")


def infer_timing_target(card_type, name, effect):
    effect = clean(effect)
    name = clean(name)

    if "即将成交" in effect:
        return "成交前", "当前拍品"
    if "暗标" in effect:
        return "暗标拍卖时", "本次暗标"
    if "荷兰" in effect:
        return "荷兰式拍卖时", "自己"
    if "竞拍开始前" in effect:
        return "竞拍开始前", "按效果指定"
    if "本日一件在拍藏品" in effect or "在拍藏品" in effect:
        return "预展后，竞拍前", "1 件在拍藏品"
    if "秘密委托" in effect:
        return "自由阶段", "1 名对手"
    if "已成交" in effect:
        return "成交后或自由阶段", "1 件已成交藏品"
    if "下一天预展" in effect:
        return "竞拍后事件窗口 / 日末", "下一天预展藏品"
    if "现金总数排名" in effect:
        return "自由阶段", "全场玩家"
    if "出售藏品给银行" in effect or "出售给银行" in effect or "卖给银行" in effect:
        return "自由阶段", "按效果指定"
    if "出价" in effect or "竞拍" in effect:
        return "竞拍中", "按效果指定"
    if "锦囊" in effect and card_type == "效果反制类":
        return "他人使用锦囊/事件时", "该效果"
    if "选择一名" in effect or "选择一件" in effect:
        return "锦囊阶段", "按效果指定"
    if not name or not effect:
        return "待补充", "待补充"
    return "按效果文本执行", "按效果文本指定"


def collect_cards():
    doc = Document(SRC)
    cards = []
    current_type = None

    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = clean(block.text)
            for keyword, mapped_type in SECTION_TYPE_KEYWORDS:
                if keyword in text:
                    current_type = mapped_type
            continue

        rows = [[clean(cell.text) for cell in row.cells] for row in block.rows]
        if not rows:
            continue

        header = rows[0]
        header_joined = "|".join(header)

        # Skip the ratio/description table.
        if "设计说明" in header_joined and "数量" in header_joined:
            continue

        # Existing full-format table.
        if "编号" in header_joined and "卡名" in header_joined and "使用时机" in header_joined:
            for row in rows[1:]:
                if len(row) < 6 or not row[0]:
                    continue
                cards.append(
                    {
                        "source_id": row[0],
                        "name": row[1],
                        "type": normalize_type(row[2]),
                        "timing": row[3],
                        "target": row[4],
                        "effect": row[5],
                        "note": "保留自原总表",
                    }
                )
            continue

        # Three-column user-added tables, sometimes with or without a header.
        for row in rows:
            if not any(row):
                continue
            source_id = row[0] if len(row) > 0 else ""
            name = row[1] if len(row) > 1 else ""
            effect = row[2] if len(row) > 2 else ""
            if source_id in {"#", "编号"} or name in {"锦囊名", "卡名"}:
                continue
            timing, target = infer_timing_target(current_type or "待分类", name, effect)
            note = "来自分组表"
            if not name or not effect:
                name = name or "待补充"
                effect = effect or "待补充"
                note = "原文只有编号或内容不完整，需补齐"
            cards.append(
                {
                    "source_id": source_id,
                    "name": name,
                    "type": current_type or "待分类",
                    "timing": timing,
                    "target": target,
                    "effect": effect,
                    "note": note,
                }
            )

    return cards


def build_doc(cards):
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    run = title.add_run("锦囊卡")
    set_font(run, size=22, bold=True, color="1F4E79")
    subtitle = doc.add_paragraph()
    run = subtitle.add_run("合并整理版：完整锦囊总表")
    set_font(run, size=11, color="666666")

    doc.add_heading("统一规则", level=2)
    for line in [
        "锦囊卡是个人战术牌，主要影响当前选择、当前拍品、当前交易或一次反制。",
        "锦囊开局每人 2 张，黑市价 30 银元/张，手牌无上限；建议每天最多使用 2 张。",
        "多人同时使用锦囊时，从当前主持人开始顺时针依次结算；同一锦囊效果最多被反制 1 次。",
        "本表合并了原总表和你新增的分组表；重复/相近效果暂不删除，只统一格式。",
    ]:
        add_bullet(doc, line)

    type_counts = {}
    for card in cards:
        type_counts[card["type"]] = type_counts.get(card["type"], 0) + 1
    count_text = "；".join(f"{key} {value} 张" for key, value in type_counts.items())
    p = doc.add_paragraph()
    run = p.add_run(f"当前合并：{len(cards)} 条。分类统计：{count_text}。")
    set_font(run, size=9.5, color="666666")

    doc.add_heading("完整锦囊表", level=2)
    rows = []
    for idx, card in enumerate(cards, start=1):
        rows.append(
            (
                str(idx),
                card["source_id"],
                card["name"],
                card["type"],
                card["timing"],
                card["target"],
                card["effect"],
                card["note"],
            )
        )

    add_table_custom(
        doc,
        ["序号", "原编号", "锦囊名", "类型", "使用时机", "目标", "效果", "备注"],
        rows,
        [0.38, 0.55, 0.9, 0.85, 1.15, 1.05, 3.35, 1.5],
        font_size=7.3,
    )

    doc.save(OUT)
    print(OUT)
    print(f"cards={len(cards)}")
    print(type_counts)


def main():
    cards = collect_cards()
    build_doc(cards)


if __name__ == "__main__":
    main()
