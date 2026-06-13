from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from organize_event_cards import (
    add_bullet,
    apply_table_widths,
    configure_document,
    set_cell_shading,
    set_cell_text,
    set_font,
    set_repeat_header,
    set_table_borders,
)


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "委托.docx"


COMMISSIONS = [
    # 收藏路线
    ("W01", "收藏大家", "收藏", "终局时持有藏品数量 >= 4。", "8", "final_items_count >= 4", "基础收藏目标，容易理解但会占用现金。"),
    ("W02", "少而精", "收藏", "终局时持有藏品数量 <= 2，且其中至少 1 件终局价值 >= 180。", "10", "final_items_count <= 2; owned_item.max(final_value) >= 180", "鼓励精品路线，避免单纯摆烂少买。"),
    ("W03", "百类杂家", "收藏", "终局时持有至少 4 种不同类别的藏品。", "10", "count_distinct(owned_item.category) >= 4", "推动横向收集和交易。"),
    ("W04", "专攻一门", "收藏", "终局时持有同一类别藏品 >= 3 件。", "10", "max_count_by(owned_item.category) >= 3", "和百类杂家互斥，抽 2 选 1 时有路线选择。"),
    ("W05", "怪谈藏家", "收藏", "终局时持有【奇物/灵器/邪物/遗物/绝笔】中任意 2 件。", "9", "owned_item.category in special_categories count >= 2", "主题感强，目标清楚，不依赖隐藏信息。"),
    # 价值与风险路线
    ("W06", "捡漏王", "价值", "终局时至少 1 件藏品满足：成交价 < 传闻下限，且终局价值 > 成交价。", "11", "owned_item.any(purchase_price < rumor_min and final_value > purchase_price)", "有爽点但要真买到便宜货，难度较高。"),
    ("W07", "高价不悔", "价值", "终局时至少 1 件藏品满足：成交价 >= 180，且终局价值 >= 成交价。", "10", "owned_item.any(purchase_price >= 180 and final_value >= purchase_price)", "奖励敢压重注但不鼓励乱烧钱。"),
    ("W08", "破烂淘金", "价值", "终局时持有至少 1 件带负面属性或赝品判定的藏品，且你的藏品总价值 >= 300。", "11", "owned_item.any(is_fake or has_negative_attr); collection_final_value >= 300", "把坏运气转成挑战，不会白送。"),
    ("W09", "真金不怕火", "价值", "终局时持有藏品 >= 3 件，且没有任何赝品藏品。", "9", "final_items_count >= 3; owned_item.fake_count == 0", "稳健路线，奖励适中。"),
    ("W10", "假作真时", "价值", "终局时持有至少 1 件赝品或【仿品】属性藏品。", "8", "owned_item.any(is_fake or attr == '仿品')", "给风险品一点补偿，但奖励不高。"),
    # 竞拍路线
    ("W11", "百发百中", "竞拍", "整局通过正式竞拍获得藏品 >= 3 件，不含玩家交易和黑市获得。", "10", "auction_win_count >= 3", "主线目标，适合主动竞价玩家。"),
    ("W12", "险胜专家", "竞拍", "至少 2 次以不超过第二高价 20 银元的差额赢得拍品。", "11", "close_win_count(diff_to_second <= 20) >= 2", "制造压线胜利的兴奋感，代码只需记录第二高价。"),
    ("W13", "暗标高手", "竞拍", "至少赢得 1 件暗标拍卖藏品。", "9", "auction_win_by_mode['暗标'] >= 1", "模式随机时有不确定性，所以奖励不过高。"),
    ("W14", "荷兰急先锋", "竞拍", "至少赢得 1 件荷兰式拍卖藏品。", "9", "auction_win_by_mode['荷兰'] >= 1", "鼓励果断喊停，条件清楚。"),
    ("W15", "打包赌徒", "竞拍", "至少赢得 1 次打包拍卖，并在终局仍持有包内至少 1 件藏品。", "10", "package_win_count >= 1; owned_item.from_won_package >= 1", "打包有盲盒感，保留到终局才算完成。"),
    # 资金路线
    ("W16", "守财奴", "资金", "终局还清贷款后，剩余现金 >= 250。", "9", "final_cash_after_debt >= 250", "现金路线可行，但会牺牲收藏。"),
    ("W17", "弹尽粮绝", "资金", "终局还清贷款后现金 <= 50，且藏品总价值 >= 350。", "10", "final_cash_after_debt <= 50; collection_final_value >= 350", "鼓励把钱变成资产，避免故意清空现金。"),
    ("W18", "银行常客", "资金", "整局贷款次数 >= 2，且终局全部还清，没有触发强制清算。", "9", "loans_taken >= 2; unpaid_loans == 0; forced_liquidation_count == 0", "贷款变成节奏工具，而不是纯惩罚。"),
    ("W19", "无债一身轻", "资金", "整局没有贷款，且终局持有藏品 >= 3 件。", "9", "loans_taken == 0; final_items_count >= 3", "给稳健玩家明确目标。"),
    ("W20", "周转大师", "资金", "整局卖给银行的藏品 >= 2 件，且终局仍持有藏品 >= 2 件。", "8", "sell_to_bank_count >= 2; final_items_count >= 2", "鼓励资产周转，但条件相对容易。"),
    # 交易与黑市路线
    ("W21", "中间商", "交易", "整局参与玩家间交易 >= 2 次，买入或卖出都计入。", "9", "player_trade_count >= 2", "推动谈判和桌面互动。"),
    ("W22", "倒手生财", "交易", "至少 1 件藏品经你买入后再卖出，且卖出价 > 买入价。", "11", "flipped_item_profit_count >= 1", "强互动目标，需要记录物品持有和交易价格。"),
    ("W23", "黑市常客", "黑市", "整局在黑市购买牌总数 >= 3。", "8", "black_market_cards_bought >= 3", "容易完成，奖励较低。"),
    ("W24", "囤牌人", "黑市", "终局时手牌数量（锦囊+事件）>= 4。", "8", "final_hand_cards_count >= 4", "鼓励保留资源，但会牺牲使用机会。"),
    ("W25", "末班车", "黑市", "第 9 天黑市至少购买 1 张牌，并在第 9 天或第 10 天使用其中 1 张。", "10", "bought_card_ids_on_day9; used_card_id in bought_card_ids_on_day9 after_day >= 9", "最后补给变成戏剧性翻盘点。"),
    # 信息与牌效路线
    ("W26", "情报达人", "信息", "整局使用信息探查类锦囊 >= 3 张。", "9", "trick_used_by_type['信息探查类'] >= 3", "直接支持信息路线。"),
    ("W27", "未卜先知", "信息", "曾提前查看牌库或未来藏品信息，并在之后拍得其中至少 1 件。", "11", "peeked_future_item_ids intersect later_won_item_ids not empty", "非常有故事性，代码用 item_id 记录即可。"),
    ("W28", "不动声色", "信息", "终局时至少 1 件你持有的藏品从未被任何信息效果探查过，且终局价值 >= 150。", "10", "owned_item.any(info_peek_count == 0 and final_value >= 150)", "奖励盲拍成功，和信息路线形成反差。"),
    ("W29", "牌技多面手", "信息", "整局至少使用过 3 种不同类型的锦囊。", "10", "count_distinct(used_trick.type) >= 3", "鼓励灵活使用，不只囤牌。"),
    ("W30", "风暴推手", "事件", "整局主动使用事件卡 >= 2 张。", "10", "event_cards_played >= 2", "事件是全局环境，奖励主动搅动局势。"),
    # 主持路线
    ("W31", "金牌主持", "主持", "你担任主持人时，成功卖出的藏品数量 >= 3。", "10", "host_successful_sales_count >= 3", "契合主持人路线，目标清晰。"),
    ("W32", "佣金达人", "主持", "整局主持佣金总收入 >= 120 银元。", "10", "total_commission_earned >= 120", "鼓励讲价和推高成交，不直接奖励自吞。"),
    ("W33", "冷槌大师", "主持", "你担任主持人时至少自吞 1 件流拍藏品，且该藏品终局价值 > 你的自吞成本。", "11", "host_pass_in_profit_item_count >= 1", "让流拍自吞成为赌博选择，不是无脑收益。"),
    ("W34", "满堂彩", "主持", "你主持过的藏品中，至少 4 件成功成交。", "11", "hosted_item_sold_count >= 4", "适合 5 人局或主持次数多的玩家，难度偏高。"),
    ("W35", "高声望拍卖官", "主持", "你主持的至少 1 件藏品成交价 >= 其传闻上限的 80%。", "10", "hosted_item.any(sold_price >= rumor_max * 0.8)", "奖励说服和气氛，不要求主持人选择拍卖模式。"),
    # 混合与逆袭路线
    ("W36", "后发制人", "逆袭", "第 7 天及以后通过竞拍获得藏品 >= 2 件。", "9", "auction_wins_on_day(day >= 7) >= 2", "冲刺期目标，适合前期低调玩家。"),
    ("W37", "稳健投资者", "逆袭", "整局没有触发强制清算，且终局藏品总价值 >= 400。", "10", "forced_liquidation_count == 0; collection_final_value >= 400", "奖励稳定经营。"),
    ("W38", "以少胜多", "逆袭", "终局持有藏品数量 <= 3，且平均藏品终局价值 >= 130。", "10", "final_items_count <= 3; avg(owned_item.final_value) >= 130", "精品路线的另一种表达。"),
    ("W39", "全场焦点", "综合", "整局至少完成：赢得 1 次竞拍、完成 1 次玩家交易、使用 1 张锦囊、主动使用 1 张事件卡。", "12", "auction_win_count >= 1; player_trade_count >= 1; trick_cards_played >= 1; event_cards_played >= 1", "跨系统目标，难但好玩，奖励最高。"),
    ("W40", "三市行商", "黑市", "第 3、6、9 天黑市都至少购买过 1 张牌。", "10", "black_market_purchase_days contains 3,6,9", "让三次黑市成为长期规划。"),
]


IMPLEMENTATION_FIELDS = [
    ("玩家终局状态", "final_cash_after_debt, final_items_count, collection_final_value, final_hand_cards_count"),
    ("藏品记录", "item_id, category, rumor_min, rumor_max, purchase_price, final_value, is_fake, attributes"),
    ("拍卖记录", "auction_mode, winner_id, sold_price, second_highest_bid, host_id, package_id"),
    ("主持记录", "host_successful_sales_count, commission_earned, pass_in/self_buy item_id 与成本"),
    ("交易记录", "buyer_id, seller_id, item_id, price, day, from_player_trade / sell_to_bank"),
    ("牌效记录", "used_card_id, card_type, target_item_id, day, peeked_item_ids, event_cards_played"),
    ("黑市记录", "day, bought_card_ids, card_count, later_used_card_ids"),
    ("贷款/清算", "loans_taken, unpaid_loans, forced_liquidation_count"),
]


def add_table_custom(doc, headers, rows, widths, font_size=7.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table, color="B7C3D0")
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, header in enumerate(headers):
        set_cell_text(hdr.cells[i], header, size=8.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr.cells[i], "385D6B")
        for run in hdr.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 4) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[i], value, size=font_size, align=align)
        route = row_data[2] if len(row_data) > 2 else ""
        if route in {"收藏", "价值"}:
            set_cell_shading(row.cells[2], "EAF2F8")
        elif route in {"竞拍", "主持"}:
            set_cell_shading(row.cells[2], "FDEBD0")
        elif route in {"资金", "交易", "黑市"}:
            set_cell_shading(row.cells[2], "E8F6EF")
        elif route in {"信息", "事件"}:
            set_cell_shading(row.cells[2], "F4ECF7")
        else:
            set_cell_shading(row.cells[2], "F2F3F4")
    apply_table_widths(table, widths)
    return table


def add_small_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table, color="B7C3D0")
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, header in enumerate(headers):
        set_cell_text(hdr.cells[i], header, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr.cells[i], "5B7F95")
        for run in hdr.cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[i], value, size=8.2, align=align)
    apply_table_widths(table, widths)
    return table


def main():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("秘密委托")
    set_font(title_run, size=22, bold=True, color="1F4E79")
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("40 张完整委托：终局目标、策略路线与代码判定字段")
    set_font(subtitle_run, size=11, color="666666")

    doc.add_heading("设计规则", level=2)
    for line in [
        "秘密委托是开局隐藏目标：每位玩家抽 2 选 1，终局揭晓，完成则获得标注声望，未完成不扣分。",
        "每名玩家终局最多结算 1 张秘密委托；若后续允许持有多张，仍只能选择 1 张完成。",
        "奖励控制在 8-12 声望：容易完成给 8-9，跨系统或高风险目标给 10-12。",
        "所有条件都采用终局或历史日志判定，避免主持人主观裁决，适合数字版实现。",
        "拍卖模式按每日系统随机生成记录，不把任何委托写成“主持人选择某拍卖模式”。",
    ]:
        add_bullet(doc, line)

    doc.add_heading("代码记录建议", level=2)
    add_small_table(
        doc,
        ["模块", "建议记录字段"],
        IMPLEMENTATION_FIELDS,
        [1.6, 8.35],
    )

    doc.add_heading("委托列表", level=2)
    p = doc.add_paragraph()
    note = p.add_run("字段说明：代码判定字段不是玩家文案，而是给数字版实现时使用的最小状态/日志依据。")
    set_font(note, size=9.2, color="666666")

    add_table_custom(
        doc,
        ["编号", "委托名称", "路线", "完成条件", "声望", "代码判定字段", "趣味/平衡说明"],
        COMMISSIONS,
        [0.42, 0.88, 0.62, 2.35, 0.42, 2.45, 2.86],
        font_size=7.15,
    )

    doc.add_heading("平衡检查", level=2)
    for line in [
        "收藏类委托分为多件、少件、多类别、单类别和主题类别，互相拉扯，能形成抽 2 选 1 的真实选择。",
        "资金类委托都附带资产、还款或不清算约束，避免玩家为了完成委托故意把现金花光或乱贷款。",
        "拍卖模式相关委托奖励不高，并且只要求赢 1 次，适配每日随机模式；若数字版不保证模式出现，可降低这类委托发放权重。",
        "主持类委托奖励主持成功、佣金和有利润的流拍自吞，不奖励单纯压低成交或拖慢游戏。",
        "信息、事件、黑市委托鼓励玩家使用系统资源，但不要求使用指定卡名，后续增删卡牌也不容易坏。",
    ]:
        add_bullet(doc, line)

    doc.save(OUT)
    print(OUT)
    print(f"commissions={len(COMMISSIONS)}")


if __name__ == "__main__":
    main()
